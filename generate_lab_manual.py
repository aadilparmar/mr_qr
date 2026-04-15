#!/usr/bin/env python3
"""
QRCode Pro — Lab Manual Generator
Generates a complete 12-practical Software Engineering lab manual DOCX
for the QRCode Pro mini project.
"""

import os, sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ══════════════════════════════════════════════════════════════════════
# CONSTANTS
# ══════════════════════════════════════════════════════════════════════
STUDENT = "Aadil Parmar(92200103068)"
BATCH   = "6EC3 - A Batch"
FONT    = "Times New Roman"
BODY_SZ = Pt(12)
HEAD2_SZ = Pt(16)
HEAD3_SZ = Pt(14)
HEAD4_SZ = Pt(12)
SMALL_SZ = Pt(10)

OUT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "92200103068_EC3_Lab Manual.docx")

doc = Document()

# ══════════════════════════════════════════════════════════════════════
# PAGE SETUP  (match original: A4, tight margins)
# ══════════════════════════════════════════════════════════════════════
for section in doc.sections:
    section.page_width  = Emu(7556500)
    section.page_height = Emu(10693400)
    section.top_margin    = Cm(1.27)
    section.bottom_margin = Cm(0.5)
    section.left_margin   = Cm(0.7)
    section.right_margin  = Cm(0.5)

# ══════════════════════════════════════════════════════════════════════
# STYLE DEFAULTS
# ══════════════════════════════════════════════════════════════════════
style_normal = doc.styles['Normal']
style_normal.font.name = FONT
style_normal.font.size = BODY_SZ
style_normal.paragraph_format.space_after = Pt(2)
style_normal.paragraph_format.space_before = Pt(0)

# ══════════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════

def add_heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = HEAD2_SZ
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_heading3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = HEAD3_SZ
    run.font.bold = True
    return p

def add_heading4(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = HEAD4_SZ
    run.font.bold = True
    return p

def add_body(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_SZ
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.name = FONT
        r1.font.size = BODY_SZ
        r1.font.bold = True
        r2 = p.add_run(text)
        r2.font.name = FONT
        r2.font.size = BODY_SZ
    else:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = BODY_SZ
    return p

def add_empty():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(4)
    return p

def add_footer_line():
    """Add the student footer at bottom of practical."""
    add_empty()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(STUDENT)
    r.font.name = FONT
    r.font.size = SMALL_SZ
    p.add_run("\t")
    p2 = doc.add_paragraph()
    r2 = p2.add_run(BATCH)
    r2.font.name = FONT
    r2.font.size = SMALL_SZ

def add_page_break():
    doc.add_page_break()

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = FONT
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "1F4E79")

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


# ══════════════════════════════════════════════════════════════════════
#  UNIVERSITY HEADER TABLE  (reusable)
# ══════════════════════════════════════════════════════════════════════
def add_uni_header():
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("FACULTY OF ENGINEERING & TECHNOLOGY")
    r1.font.name = FONT
    r1.font.size = Pt(13)
    r1.font.bold = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Department of Computer Engineering")
    r2.font.name = FONT
    r2.font.size = BODY_SZ
    r2.font.bold = True

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("01CE0607 - Software Engineering \u2013 Lab Manual")
    r3.font.name = FONT
    r3.font.size = BODY_SZ
    add_empty()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                    PRACTICAL 1                                   ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_1():
    add_uni_header()
    add_heading2("Practical 1")
    add_body("Problem Statement with Purpose, Scope, Literature Review, and Future Scope", bold=True)
    add_empty()
    add_body("Aim: Identify a relevant problem or project definition. Write a detailed problem statement for the system, along with its Purpose, Scope, Existing system details with a literature review and mention Future scope of the system.")
    add_empty()

    add_heading3("QRCode Pro \u2014 Professional QR Code Generator")
    add_empty()

    # Introduction
    add_heading4("Introduction")
    add_body("QRCode Pro is a comprehensive, web-based QR code generation platform built using PHP 8.0+, SQLite, Tailwind CSS, and JavaScript. The system enables users to generate, customize, scan, and manage over 24 different types of QR codes through a single, elegant interface. It features 62 customization controls spanning dot styles, gradients, eye frames, logos, shadows, and frames \u2014 along with bulk generation capabilities for up to 100 QR codes simultaneously, an integrated camera-based scanner, and full export support for PNG, SVG, and PDF formats.")
    add_empty()

    # Problem Statement
    add_heading4("Problem Statement")
    add_body("In today\u2019s digital ecosystem, QR codes have become ubiquitous \u2014 used in payments, marketing, authentication, and information sharing. However, the existing landscape of QR code tools suffers from several critical shortcomings:")
    add_bullet("Most free QR generators support only 3\u20135 basic types (URL, text, WiFi) while ignoring specialized types such as UPI payments, vCard contacts, calendar events, cryptocurrency addresses, and social media profiles.")
    add_bullet("Customization is severely limited \u2014 users are typically restricted to basic foreground/background color changes, with no control over dot styles, eye frames, gradients, logo overlays, or shadow effects.")
    add_bullet("Bulk generation capabilities are rare or locked behind paid plans, forcing businesses to generate QR codes one at a time.")
    add_bullet("Most tools are stateless \u2014 they offer no user accounts, no generation history, no favorites, and no dashboard analytics.")
    add_bullet("QR scanning functionality is typically provided by separate applications, creating a fragmented user experience.")
    add_bullet("Privacy concerns exist as most online QR generators process data on remote servers, with no transparency about data retention.")
    add_body("There is a clear need for a unified, self-hosted, open-source QR code platform that combines comprehensive generation, deep customization, bulk processing, scanning, and history management \u2014 all completely free and privacy-respecting.")
    add_empty()

    # Purpose
    add_heading4("Purpose of the System")
    add_body("The primary purpose of QRCode Pro is to provide a single, unified platform that addresses every aspect of QR code workflow:")
    add_bullet("Generation: ", "Support 24+ QR code types including URL, WiFi, vCard, UPI, WhatsApp, Email, SMS, Calendar Events, GPS Locations, Bitcoin, YouTube, Twitter, Instagram, Facebook, LinkedIn, Spotify, Zoom, PDF, Image, MeCard, and App Store links.")
    add_bullet("Customization: ", "Offer 62 fine-grained controls across 9 categories \u2014 color presets, dot style & gradient, background (transparent/gradient/rounded), eye frame style & gradient, eye dot style & gradient, logo/watermark overlay, frame & label, shadow & effects, and size & export settings.")
    add_bullet("Bulk Processing: ", "Enable generation of up to 100 QR codes simultaneously from manual input or CSV upload, with ZIP archive download.")
    add_bullet("Scanning: ", "Provide an integrated QR scanner supporting both camera-based real-time scanning and image file upload.")
    add_bullet("History & Management: ", "Offer persistent storage of all generated QR codes with search, filtering, favorites, pagination, and CSV export.")
    add_bullet("Analytics: ", "Present dashboard statistics with Chart.js-powered visualizations showing QR type distribution and 7-day activity trends.")
    add_bullet("Privacy: ", "Ensure all data stays on the self-hosted server with bcrypt password hashing, session security, and no third-party data sharing.")
    add_empty()

    # Scope
    add_heading4("Scope of the System")
    add_body("The scope of QRCode Pro encompasses the following functional modules:")
    add_bullet("User Authentication Module: Registration with full name, username, email, and password (bcrypt hashed). Login via username or email. Session management with regeneration to prevent fixation attacks. Profile management for updating personal details and changing passwords.")
    add_bullet("QR Code Generation Module: 23 distinct QR type forms with type-specific input fields. Real-time preview using the qr-code-styling JavaScript library. 62 customization controls organized in 9 collapsible accordion sections. Export as PNG, SVG, or embedded in PDF via jsPDF.")
    add_bullet("QR Code History Module: Paginated grid view of all saved QR codes. Filtering by type and favorites. Client-side search across titles, content, and types. Inline delete, favorite toggle, copy, and recreate actions. CSV export of filtered results.")
    add_bullet("Bulk Generation Module: Manual entry (one item per line, max 100) or CSV file upload. Configurable type, size, and colors. Progress bar with thumbnail previews. ZIP archive download with manifest file.")
    add_bullet("QR Scanner Module: Camera-based scanning using html5-qrcode library with multi-camera support. Image upload scanning with drag-and-drop support. Auto-detection of QR content type (URL, email, phone, WiFi, text). Session-based scan history with copy and open actions.")
    add_bullet("Dashboard Module: Four stat cards (total QR codes, total scans, types used, current plan). Quick-generate grid for 8 popular QR types. Recent QR codes grid (last 8). Doughnut chart for QR type distribution and bar chart for 7-day activity.")
    add_bullet("Public Pages: Landing page with hero section, stats, type grid, and features. About page with team information and tech stack. Contact form with database storage. Privacy policy page.")
    add_bullet("API Module: JSON-based POST endpoint (api/generate.php) for programmatic QR code generation with type validation and auth guard.")
    add_empty()

    # Literature Review
    add_heading4("Literature Review")
    add_empty()

    lit_headers = ["Title", "Author/Tool", "Year", "Method", "Advantages", "Disadvantages", "Tools Used", "Remark"]
    lit_rows = [
        ["QR Code Generation using Python", "Jain et al.", "2021", "Python qrcode library with PIL", "Simple implementation, open-source", "No customization, single type only", "Python, qrcode, PIL", "Basic text-to-QR only"],
        ["QR Monkey (qrmonkey.com)", "QR Monkey", "2023", "Web-based generator with basic colors", "Free, multiple types, logo upload", "No bulk, no history, limited styles", "JavaScript, Canvas API", "Popular but feature-limited"],
        ["QR Code Generator Pro (qr-code-generator.com)", "Egoditor GmbH", "2022", "SaaS with freemium model", "Dynamic QR, analytics, professional", "Paid features, no self-hosting, data on their servers", "React, Node.js, Cloud", "Enterprise but not privacy-focused"],
        ["Dynamic QR Code System for Smart Cities", "Sharma & Kumar", "2022", "Server-side QR with redirect tracking", "Scan analytics, URL redirect", "Complex setup, server-dependent", "PHP, MySQL, Google Charts", "Research-focused implementation"],
        ["GoQR.me API", "GoQR.me", "2023", "REST API for QR generation", "API-first, multiple formats", "No UI, no auth, no customization beyond colors", "Java, REST API", "Developer tool only"],
        ["QR Code Styling Library", "Nickolay Bliokh", "2023", "Client-side canvas with dot/eye styling", "Deep customization, gradients, logos", "Library only, no full application", "JavaScript, Canvas", "Foundation for QRCode Pro\u2019s engine"],
    ]
    make_table(lit_headers, lit_rows, [1.0, 0.8, 0.5, 1.0, 1.0, 1.0, 0.8, 0.9])
    add_empty()

    add_body("The literature review reveals that while numerous QR code tools exist, none combines comprehensive type support, deep visual customization, bulk generation, integrated scanning, and user history management in a single, free, self-hosted, open-source package. QRCode Pro fills this gap by integrating all these capabilities into one cohesive platform.")
    add_empty()

    # Future Scope
    add_heading4("Future Scope")
    add_body("The evolution of QRCode Pro will advance with emerging technologies and growing user demands:")
    add_empty()

    add_heading4("AI-Powered QR Design")
    add_bullet("Integrate AI models to automatically suggest color schemes and dot styles based on brand logos or images uploaded by users.")
    add_bullet("Auto-generate visually aesthetic QR codes that match a company\u2019s brand identity using color extraction algorithms.")
    add_bullet("AI-based readability scoring to warn users before they download a QR code that may be difficult to scan.")
    add_empty()

    add_heading4("Dynamic QR Codes with Scan Analytics")
    add_bullet("Implement redirect-based dynamic QR codes where the destination URL can be changed after printing.")
    add_bullet("Track scan counts, geographic locations, device types, and timestamps for each QR code.")
    add_bullet("Dashboard analytics with heatmaps showing scan patterns over time.")
    add_empty()

    add_heading4("API Marketplace & Integrations")
    add_bullet("Develop a full REST API with API key authentication for developers to generate QR codes programmatically.")
    add_bullet("Create WordPress, Shopify, and WooCommerce plugins for seamless e-commerce integration.")
    add_bullet("Webhook support to notify external systems when QR codes are scanned.")
    add_empty()

    add_heading4("Progressive Web App & Mobile")
    add_bullet("Convert QRCode Pro into a Progressive Web App (PWA) with offline scanning capabilities.")
    add_bullet("Develop native Android and iOS apps using React Native or Flutter for enhanced camera scanning performance.")
    add_bullet("Push notifications for scan alerts and QR code expiry reminders.")
    add_empty()

    add_heading4("Enterprise Features")
    add_bullet("Team workspaces with role-based access control (admin, editor, viewer) for collaborative QR management.")
    add_bullet("White-label support allowing businesses to deploy QRCode Pro under their own branding.")
    add_bullet("SAML/SSO integration for enterprise single sign-on authentication.")

    add_footer_line()
    add_page_break()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                    PRACTICAL 2                                   ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_2():
    add_uni_header()
    add_heading2("Practical 2")
    add_heading3("Software Development Life Cycle (SDLC) & Project Planning")
    add_body("Aim: Study and apply SDLC. Select a process model and create a project plan for the selected system using Gantt chart. Include a work breakdown structure and a basic timeline for development.")
    add_empty()

    add_heading4("Software Development Life Cycle")
    add_body("The Software Development Life Cycle (SDLC) is a systematic process for planning, creating, testing, and deploying a software system. For QRCode Pro, we adopt the Waterfall Model due to its structured, sequential approach which is well-suited for a university mini project with clearly defined requirements and a fixed timeline.")
    add_empty()

    add_heading4("Why Waterfall Model for QRCode Pro")
    add_bullet("Requirements are well-defined upfront \u2014 24+ QR types, customization controls, bulk generation, scanner, and history management.")
    add_bullet("The project has a fixed deadline (semester-end submission) making iterative development impractical.")
    add_bullet("The team size is small (3 members) and all team members understand the full scope from the beginning.")
    add_bullet("Documentation is a primary deliverable (lab manual, report), which aligns with Waterfall\u2019s documentation-heavy approach.")
    add_empty()

    add_heading4("Waterfall Phases for QRCode Pro")
    add_empty()

    add_heading4("Phase 1: Requirements Gathering (Week 1\u20132)")
    add_bullet("Analyzed existing QR code tools (QR Monkey, GoQR.me, QR Code Generator Pro) to identify feature gaps.")
    add_bullet("Defined 24+ QR code types with their encoding formats (URL, WiFi SSID, vCard, UPI, etc.).")
    add_bullet("Documented 62 customization controls across 9 categories.")
    add_bullet("Established non-functional requirements: real-time preview (<200ms), bulk generation (100 codes), responsive design, dark mode support.")
    add_empty()

    add_heading4("Phase 2: System Design (Week 3\u20134)")
    add_bullet("Designed database schema with 4 tables: users, qr_codes, bulk_jobs, contact_messages.")
    add_bullet("Created system architecture: 3-tier (Presentation, Business Logic, Data) with PHP backend and SQLite storage.")
    add_bullet("Designed clean URL routing system with .htaccess (Apache) and router.php (built-in server) support.")
    add_bullet("Planned UI/UX with dual design system: landing pages (indigo brand) and app pages (orange accent).")
    add_empty()

    add_heading4("Phase 3: Implementation (Week 5\u20139)")
    add_bullet("Week 5\u20136: Core infrastructure \u2014 database config, auth system, header/footer, routing.")
    add_bullet("Week 6\u20137: QR generation engine \u2014 23 type forms, qr-code-styling integration, real-time preview.")
    add_bullet("Week 7\u20138: Customization engine \u2014 62 controls in 9 accordion sections, gradient support, logo overlay.")
    add_bullet("Week 8\u20139: Supporting features \u2014 bulk generator, QR scanner, dashboard analytics, history management.")
    add_empty()

    add_heading4("Phase 4: Testing (Week 10\u201311)")
    add_bullet("Unit testing of all 23 QR type encoders in encodeQRContent() function.")
    add_bullet("Integration testing of auth flow (register \u2192 login \u2192 generate \u2192 save \u2192 history).")
    add_bullet("Cross-browser testing on Chrome, Firefox, Safari, and Edge.")
    add_bullet("Mobile responsiveness testing with bottom navigation bar.")
    add_bullet("Dark mode testing across all pages.")
    add_empty()

    add_heading4("Phase 5: Deployment & Maintenance (Week 12)")
    add_bullet("Deployed on Apache server with .htaccess configuration for clean URLs and security.")
    add_bullet("Database initialized automatically via initDatabase() on first request.")
    add_bullet("Documentation: lab manual, project report, and PowerPoint presentation.")
    add_empty()

    # Gantt Chart as Table
    add_heading4("Gantt Chart")
    gantt_headers = ["Phase / Task", "W1", "W2", "W3", "W4", "W5", "W6", "W7", "W8", "W9", "W10", "W11", "W12"]
    gantt_rows = [
        ["Requirements Gathering",     "\u2588", "\u2588", "",   "",   "",   "",   "",   "",   "",   "",    "",    ""],
        ["  \u2022 Market Analysis",    "\u2588", "",   "",   "",   "",   "",   "",   "",   "",   "",    "",    ""],
        ["  \u2022 Feature Definition", "",   "\u2588", "",   "",   "",   "",   "",   "",   "",   "",    "",    ""],
        ["System Design",              "",   "",   "\u2588", "\u2588", "",   "",   "",   "",   "",   "",    "",    ""],
        ["  \u2022 Database Schema",    "",   "",   "\u2588", "",   "",   "",   "",   "",   "",   "",    "",    ""],
        ["  \u2022 Architecture & UI",  "",   "",   "",   "\u2588", "",   "",   "",   "",   "",   "",    "",    ""],
        ["Implementation",             "",   "",   "",   "",   "\u2588", "\u2588", "\u2588", "\u2588", "\u2588", "",    "",    ""],
        ["  \u2022 Core + Auth",        "",   "",   "",   "",   "\u2588", "\u2588", "",   "",   "",   "",    "",    ""],
        ["  \u2022 QR Engine (62 ctrl)","",   "",   "",   "",   "",   "\u2588", "\u2588", "",   "",   "",    "",    ""],
        ["  \u2022 Bulk + Scanner",     "",   "",   "",   "",   "",   "",   "",   "\u2588", "\u2588", "",    "",    ""],
        ["Testing",                     "",   "",   "",   "",   "",   "",   "",   "",   "",   "\u2588", "\u2588", ""],
        ["Deployment & Docs",           "",   "",   "",   "",   "",   "",   "",   "",   "",   "",    "",    "\u2588"],
    ]
    make_table(gantt_headers, gantt_rows)
    add_empty()

    # Team
    add_heading4("Team Structure & Roles")
    team_headers = ["Role", "Member", "Enrollment", "Responsibilities"]
    team_rows = [
        ["Lead Developer", "Aadil Parmar", "92200103068", "Architecture, QR engine, 62 customization controls, API design"],
        ["Developer", "Dhruvil Janani", "92420103002", "Auth system, dashboard analytics, history management, bulk generator"],
        ["Developer", "Dhaval Chauhan", "92100103297", "QR scanner, contact form, public pages (about, privacy), testing"],
    ]
    make_table(team_headers, team_rows, [1.2, 1.2, 1.2, 3.4])
    add_empty()

    # Risks
    add_heading4("Risks and Mitigation")
    add_body("Every software project carries risks. For QRCode Pro, identified risks and mitigations include:")
    add_bullet("Browser Compatibility: ", "Client-side QR rendering may behave differently across browsers. Mitigated by using the well-tested qr-code-styling library and cross-browser testing.")
    add_bullet("Performance with 62 Controls: ", "Real-time QR preview with many customization inputs could cause lag. Mitigated by implementing 160ms debounced updates in the updateQR() function.")
    add_bullet("SQLite Concurrency: ", "SQLite has limited concurrent write support. Mitigated by enabling WAL journal mode and keeping transactions short.")
    add_bullet("Scope Creep: ", "The temptation to add more QR types or features. Mitigated by freezing requirements after Week 2.")
    add_empty()

    # Deliverables
    add_heading4("Deliverables")
    add_bullet("Requirements Document: Functional and non-functional requirements for all 8 modules.")
    add_bullet("System Design Artifacts: ER diagrams, DFDs, class diagrams, sequence diagrams, and architecture diagrams.")
    add_bullet("Functional Application: Fully deployed QRCode Pro with 24+ QR types and 62 controls.")
    add_bullet("User Documentation: About page with team info, privacy policy, and in-app help tips.")
    add_bullet("Testing Reports: Results of unit, integration, cross-browser, and mobile testing.")
    add_body("This structured planning ensures QRCode Pro is developed systematically following the Waterfall Model, with clear timelines and proactive risk mitigation.", italic=True)

    add_footer_line()
    add_page_break()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                    PRACTICAL 3                                   ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_3():
    add_uni_header()
    add_heading2("Practical-3")
    add_body("Cost and Effort Estimation for Software Development", bold=True)
    add_body("Aim: Perform a cost and effort estimation for the selected system by understanding the scope of the software to be developed.")
    add_empty()

    add_body("Cost estimation and effort analysis are essential for planning and managing software projects. Function Point Analysis (FPA) is used to estimate software size, effort, and cost based on the functionality delivered to users.")
    add_empty()

    add_heading3("Cost and Effort Estimation")
    add_body("To estimate the cost and effort for developing QRCode Pro using Function Point Analysis (FPA), the following steps are utilized.")
    add_empty()

    add_heading4("3.1 Understand the Scope of the System")
    add_body("QRCode Pro includes the following modules:")
    add_bullet("QR Generation Module: Accepts user input for 23 QR types and generates customized QR codes with 62 controls.")
    add_bullet("Authentication Module: User registration, login, logout, session management, and profile management.")
    add_bullet("History & Favorites Module: Persistent storage, search, filter, paginate, delete, and export QR codes.")
    add_bullet("Bulk Generation Module: Generate up to 100 QR codes from manual input or CSV, download as ZIP.")
    add_bullet("Scanner Module: Camera-based and file-upload QR code scanning with content detection.")
    add_bullet("Dashboard & Analytics: Stats cards, quick-generate shortcuts, Chart.js visualizations.")
    add_bullet("Public Pages: Landing page, about, contact (with DB storage), privacy policy.")
    add_bullet("REST API: JSON endpoint for programmatic QR generation with auth guard.")
    add_empty()

    add_heading4("3.2 Identify and Classify Function Types")
    add_body("Function Point Analysis for QRCode Pro classifies system components into five elements:")
    add_bullet("External Inputs (EI): ", "User inputs that modify internal data \u2014 registration form, login form, QR generation (23 type forms), profile update, password change, contact form, bulk data input, CSV upload, logo upload, favorite toggle, QR delete.")
    add_bullet("External Outputs (EO): ", "Data outputs presented to users \u2014 QR code preview (real-time canvas), PNG download, SVG download, PDF export, print output, clipboard copy, dashboard charts (doughnut + bar), history grid with QR previews, CSV export, flash messages.")
    add_bullet("External Inquiries (EQ): ", "Read-only queries \u2014 history search/filter, type selector, camera scanner decode, file upload decode, dashboard stats query, profile view, scan history display.")
    add_bullet("Internal Logical Files (ILF): ", "Internal data stores \u2014 users table, qr_codes table, bulk_jobs table, contact_messages table.")
    add_bullet("External Interface Files (EIF): ", "External data referenced \u2014 qr-code-styling CDN library, Chart.js library, html5-qrcode library, QRCode.js library, Tailwind CSS CDN, Google Fonts, Lucide Icons, JSZip, FileSaver.js, jsPDF.")
    add_empty()

    add_heading4("Table 3.1: Functionality Breakdown")
    fp_headers = ["Function Type", "Component", "Complexity"]
    fp_rows = [
        ["EI", "Registration form (4 fields + validation)", "Average"],
        ["EI", "Login form (username/email + password)", "Low"],
        ["EI", "QR Generation (23 type-specific forms)", "High"],
        ["EI", "62 Customization controls (colors, gradients, logo, frame, shadow)", "High"],
        ["EI", "Profile update (name, email)", "Low"],
        ["EI", "Password change (current + new + confirm)", "Average"],
        ["EI", "Contact form (name, email, subject, message)", "Average"],
        ["EI", "Bulk input (manual text or CSV upload)", "Average"],
        ["EI", "Logo file upload", "Low"],
        ["EI", "Favorite toggle / QR delete", "Low"],
        ["EO", "Real-time QR preview (qr-code-styling canvas)", "High"],
        ["EO", "PNG/SVG download with frame/shadow compositing", "High"],
        ["EO", "PDF export via jsPDF", "Average"],
        ["EO", "Dashboard analytics (2 Chart.js charts)", "Average"],
        ["EO", "History grid with mini QR previews", "Average"],
        ["EO", "CSV export of history", "Low"],
        ["EO", "Bulk ZIP download with manifest", "Average"],
        ["EQ", "History search & filter", "Average"],
        ["EQ", "QR type selector (23 types)", "Low"],
        ["EQ", "Camera QR scanner decode", "High"],
        ["EQ", "File upload QR decode", "Average"],
        ["EQ", "Dashboard stats aggregation", "Average"],
        ["ILF", "users table (9 columns)", "Low"],
        ["ILF", "qr_codes table (10 columns with JSON settings)", "Average"],
        ["ILF", "bulk_jobs table (6 columns)", "Low"],
        ["ILF", "contact_messages table (7 columns)", "Low"],
        ["EIF", "qr-code-styling library", "Average"],
        ["EIF", "Chart.js, html5-qrcode, QRCode.js", "Average"],
        ["EIF", "Tailwind CSS, Google Fonts, Lucide Icons", "Low"],
        ["EIF", "JSZip, FileSaver.js, jsPDF", "Low"],
    ]
    make_table(fp_headers, fp_rows, [0.8, 4.0, 1.0])
    add_empty()

    add_heading4("Table 3.2: Assigning Weights for Functions")
    w_headers = ["Function Type", "Low", "Average", "High"]
    w_rows = [
        ["External Inputs (EI)", "3", "4", "6"],
        ["External Outputs (EO)", "4", "5", "7"],
        ["External Inquiries (EQ)", "3", "4", "6"],
        ["Internal Logical Files (ILF)", "7", "10", "15"],
        ["External Interface Files (EIF)", "5", "7", "10"],
    ]
    make_table(w_headers, w_rows, [2.0, 1.0, 1.0, 1.0])
    add_empty()

    add_heading4("Calculating Unadjusted Function Points (UFP)")
    add_body("UFP = \u2211 (Function Count \u00d7 Weight)")
    add_empty()

    add_heading4("Table 3.3: Function Point Calculation")
    calc_headers = ["Function Type", "Low (Count\u00d7Wt)", "Average (Count\u00d7Wt)", "High (Count\u00d7Wt)", "Subtotal"]
    calc_rows = [
        ["EI",  "4\u00d73=12", "4\u00d74=16", "2\u00d76=12", "40"],
        ["EO",  "1\u00d74=4",  "4\u00d75=20", "2\u00d77=14", "38"],
        ["EQ",  "1\u00d73=3",  "3\u00d74=12", "1\u00d76=6",  "21"],
        ["ILF", "3\u00d77=21", "1\u00d710=10","0",           "31"],
        ["EIF", "2\u00d75=10", "2\u00d77=14", "0",           "24"],
        ["", "", "", "Total UFP", "154"],
    ]
    make_table(calc_headers, calc_rows, [1.2, 1.4, 1.4, 1.4, 0.8])
    add_empty()

    # VAF
    add_heading4("Adjust for Complexity Factors")
    add_body("The Value Adjustment Factor (VAF) accounts for 14 General System Characteristics (GSCs), each rated 0\u20135:")
    add_body("VAF = 0.65 + (TDI \u00d7 0.01)", bold=True)
    add_empty()

    add_heading4("Table 3.4: General System Characteristics for QRCode Pro")
    gsc_headers = ["#", "Characteristic", "Rating", "Justification"]
    gsc_rows = [
        ["1",  "Data Communications",       "4", "REST API, CDN integrations, AJAX calls"],
        ["2",  "Distributed Data Processing","2", "Single-server SQLite, client-side rendering"],
        ["3",  "Performance",               "4", "Real-time preview, 160ms debounce, instant scanning"],
        ["4",  "Heavily Used Configuration", "3", "62 customization controls, 9 accordion sections"],
        ["5",  "Transaction Rate",          "3", "Moderate \u2014 QR save, bulk generate, scan operations"],
        ["6",  "Online Data Entry",         "5", "23 type-specific forms, 62 controls, bulk input, CSV upload"],
        ["7",  "End-User Efficiency",       "5", "Live preview, color presets, drag-drop, keyboard shortcuts"],
        ["8",  "Online Update",             "4", "Profile update, favorites, delete, password change"],
        ["9",  "Complex Processing",        "4", "QR encoding (23 formats), canvas compositing, gradient math"],
        ["10", "Reusability",               "3", "Shared functions.php, reusable header/footer components"],
        ["11", "Installation Ease",         "4", "Zero-config SQLite, auto-init schema, .htaccess routing"],
        ["12", "Operational Ease",          "4", "Flash messages, intuitive UI, dark mode, mobile nav"],
        ["13", "Multiple Sites",            "3", "Works on Apache (subdir) and PHP built-in server (root)"],
        ["14", "Facilitate Change",         "4", "Modular PHP includes, CSS variables, configurable base path"],
    ]
    make_table(gsc_headers, gsc_rows, [0.3, 1.8, 0.5, 3.2])
    add_empty()

    add_body("TDI (Total Degree of Influence) = 4+2+4+3+3+5+5+4+4+3+4+4+3+4 = 52")
    add_body("VAF = 0.65 + (52 \u00d7 0.01) = 0.65 + 0.52 = 1.17")
    add_empty()

    add_heading4("AFP")
    add_body("AFP = UFP \u00d7 VAF")
    add_body("= 154 \u00d7 1.17")
    add_body("= 180.18 \u2248 180", bold=True)
    add_empty()

    add_heading4("Effort Estimation")
    add_body("Effort is estimated based on AFP and productivity rate.")
    add_body("For QRCode Pro, a Productivity Rate of 20 FP/Person-Month is assumed based on:")
    add_bullet("Moderately experienced student team")
    add_bullet("Standard development tools (VS Code, XAMPP, Git)")
    add_bullet("Moderate algorithmic complexity (QR encoding + canvas compositing)")
    add_empty()
    add_body("Effort (Person-Months) = AFP / Productivity Rate = 180 / 20 = 9 Person-Months", bold=True)
    add_body("With 3 team members: 9 / 3 = 3 months calendar time", bold=True)
    add_empty()

    add_heading4("Cost Estimation")
    add_body("Assumed Developer Monthly Cost = \u20b930,000")
    add_body("Cost = Effort \u00d7 Monthly Cost = 9 \u00d7 30,000")
    add_body("Total Cost = \u20b92,70,000", bold=True)
    add_body("(Note: Cost and effort values are approximate and based on assumed parameters.)", italic=True)

    add_footer_line()
    add_page_break()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                    PRACTICAL 4                                   ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_4():
    add_uni_header()
    add_heading2("Practical 4")
    add_heading3("Software Requirement Specification (SRS) Development")
    add_body("Aim: Perform a requirement analysis and develop a Software Requirement Specification (SRS) sheet for the selected system.")
    add_empty()

    add_heading3("Software Requirement Specification (SRS)")
    add_body("A Software Requirement Specification (SRS) is a structured document that defines the functionalities, constraints, and requirements of a software system. It serves as a foundation for development, testing, and stakeholder communication.")
    add_empty()

    add_heading3("Overview")
    add_body("QRCode Pro is a comprehensive, web-based QR code generation platform designed to create, customize, scan, and manage over 24 types of QR codes through a single, elegant interface. It replaces fragmented QR tools with a unified, self-hosted solution featuring 62 customization controls, bulk generation for up to 100 codes, an integrated camera/upload scanner, persistent history with analytics, and full export support (PNG, SVG, PDF).")
    add_body("The system enhances productivity for businesses, developers, and individuals by providing professional-grade QR code tools completely free and open source, with all data stored locally for maximum privacy.")
    add_body("This SRS document follows the IEEE 830-1998 standard and outlines the functional and non-functional requirements, external interfaces, constraints, and dependencies.")
    add_empty()

    add_heading3("Purpose")
    add_body("The primary purpose of QRCode Pro is to provide an automated platform for generating, customizing, and managing professional QR codes across 24+ types. The system allows users to create QR codes with 62 fine-grained visual controls, generate up to 100 codes in bulk, scan existing QR codes via camera or image upload, and maintain a searchable history of all generated codes with favorites and analytics.")
    add_empty()

    add_heading3("Document Conventions")
    add_body("This document follows IEEE 830-1998 format with numbered sections, bullet points, and standard technical terminology. It is intended for:")
    add_bullet("Developers: ", "To understand and implement system functionality and architecture.")
    add_bullet("Testers: ", "To validate system features, performance, and edge cases.")
    add_bullet("Project Guide: ", "To review system capabilities and evaluate project completeness.")
    add_bullet("Stakeholders: ", "To understand the system\u2019s scope, limitations, and design decisions.")
    add_empty()

    add_heading3("Overall Description")
    add_heading4("Product Perspective")
    add_body("QRCode Pro is a standalone, self-hosted web application that can run on any PHP 8.0+ server with SQLite support. It requires no external database server, no build tools, and no package manager \u2014 making it truly zero-configuration. The system can operate behind Apache (.htaccess routing) or PHP\u2019s built-in development server (router.php).")
    add_empty()

    add_heading4("Product Functions")
    add_bullet("QR Code Generation: Generate QR codes for 23 distinct types with type-specific input forms.")
    add_bullet("Visual Customization: 62 controls across 9 sections (colors, dots, background, eye frame, eye dot, logo, frame & label, shadow, size & export).")
    add_bullet("Bulk Generation: Process up to 100 QR codes simultaneously from text input or CSV file.")
    add_bullet("QR Scanning: Decode QR codes via camera (multi-camera support) or image upload (drag-and-drop).")
    add_bullet("History Management: Searchable, filterable, paginated history with favorites, delete, and CSV export.")
    add_bullet("Dashboard Analytics: Stat cards, quick-generate grid, Chart.js doughnut and bar charts.")
    add_bullet("User Authentication: Registration, login (username or email), profile management, password change.")
    add_bullet("Multi-format Export: PNG (with frame/shadow compositing), SVG, PDF (via jsPDF), print, and clipboard copy.")
    add_empty()

    add_heading4("Users of the System")
    add_bullet("Guest Users: ", "Can view landing page, about, contact, and privacy pages. Must register to generate QR codes.")
    add_bullet("Registered Users: ", "Full access to generation (24+ types), customization (62 controls), bulk generation, scanning, history, dashboard, and profile management.")
    add_empty()

    add_heading4("General Constraints")
    add_bullet("The system requires PHP 8.0+ with PDO and SQLite3 extensions enabled.")
    add_bullet("QR code generation is client-side (JavaScript) \u2014 the server stores metadata only.")
    add_bullet("Camera scanning requires HTTPS in production (browser security requirement).")
    add_bullet("SQLite supports limited concurrent writes (mitigated via WAL mode).")
    add_bullet("Maximum bulk generation is capped at 100 codes per batch.")
    add_empty()

    add_heading3("Specific Requirements")
    add_heading4("Functional Requirements")
    add_body("FR-01: User Registration \u2014 Accept username (min 3 chars), email (validated), password (min 6 chars, confirmed), and optional full name. Hash password with bcrypt. Regenerate session ID after registration.")
    add_body("FR-02: User Login \u2014 Accept username or email with password. Verify against bcrypt hash. Regenerate session ID on success. Redirect to dashboard.")
    add_body("FR-03: QR Code Generation \u2014 Provide 23 type-specific input forms. Encode content using encodeQRContent() with correct format for each type (tel:, mailto:, WIFI:, BEGIN:VCARD, geo:, upi://, bitcoin:, etc.).")
    add_body("FR-04: QR Customization \u2014 Provide 62 controls: 8 color presets, dot type (6 styles) + color + gradient, background color + transparent + gradient + rounded, corner square style (4) + color + gradient, corner dot style (3) + color + gradient, logo upload + size + margin + opacity + hide-dots, frame enable + color + width + padding + radius + label + label color + label size, shadow enable + color + blur + offset X/Y, QR shape + export size (4) + error correction (4) + quiet zone + format.")
    add_body("FR-05: Save to History \u2014 Save QR code metadata (type, title, content, filename, settings JSON) to qr_codes table. Redirect to history page.")
    add_body("FR-06: History Management \u2014 Display paginated grid (12 per page). Support type filter and favorites filter. Client-side search. Delete with confirmation. Favorite toggle. CSV export.")
    add_body("FR-07: Bulk Generation \u2014 Accept manual text (one per line) or CSV upload. Validate max 100 items. Generate QR codes client-side with progress bar. Download as ZIP with manifest.txt.")
    add_body("FR-08: QR Scanner \u2014 Support camera scanning (html5-qrcode, multi-camera, 10fps) and file upload (drag-drop). Auto-detect content type (URL/email/phone/WiFi/text). Maintain session scan history (max 20).")
    add_body("FR-09: Dashboard \u2014 Display 4 stat cards, 8-item quick-generate grid, last 8 QR codes, and 2 analytics charts.")
    add_body("FR-10: Contact Form \u2014 Accept name, email, subject, message (min 10 chars). Validate and store in contact_messages table.")
    add_empty()

    add_heading4("External Interfaces")
    add_body("User Interface: ", bold=True)
    add_body("Responsive web interface built with Tailwind CSS (CDN). Dual design system \u2014 landing pages use indigo brand palette, app pages use orange (#FF5C35) accent. Dark mode via html.dark class with localStorage persistence. Mobile bottom navigation bar for authenticated users. Lucide icons throughout.")
    add_body("Hardware Interface: ", bold=True)
    add_body("Camera access via WebRTC for QR scanning. No specialized hardware required. Works on desktops, laptops, tablets, and mobile phones.")
    add_body("Software Interfaces: ", bold=True)
    add_body("SQLite database via PHP PDO. External CDN libraries: Tailwind CSS, Google Fonts (Outfit, Plus Jakarta Sans, JetBrains Mono), Lucide Icons, QRCode.js, qr-code-styling, Chart.js, html5-qrcode, JSZip, FileSaver.js, jsPDF.")
    add_body("Communication Interfaces: ", bold=True)
    add_body("REST API endpoint (POST /api/generate) accepting JSON body with type, data, title, size, fg_color, bg_color. Returns JSON with success status, ID, content, and settings.")
    add_empty()

    add_heading4("Non-Functional Requirements")
    add_body("Performance: Real-time QR preview with 160ms debounce. Scanner at 10fps. Bulk generation with 50ms delay between items for UI responsiveness.")
    add_body("Security: Bcrypt password hashing. Session regeneration on login/register. CSRF token support. Input validation and htmlspecialchars() output escaping. Directory blocking via .htaccess for config/, data/, includes/.")
    add_body("Usability: Clean, modern UI with animations (fadeInUp, scaleIn). Color-coded QR type icons. Accordion-based customization to avoid overwhelming users. Mobile-first responsive design.")
    add_body("Scalability: Modular PHP includes for easy feature addition. New QR types can be added by extending getQRTypes() and encodeQRContent().")
    add_body("Availability: Zero-config SQLite with auto-initialization. No external dependencies for core functionality.")
    add_empty()

    add_heading4("Design Constraints")
    add_body("Technology Stack: PHP 8.0+ (backend), SQLite (database), Tailwind CSS via CDN (styling), Vanilla JavaScript (QR engine, UI logic).")
    add_body("Security Compliance: Passwords never stored in plain text. Session cookies are httpOnly. Sensitive directories blocked from web access.")
    add_body("Deployment Environment: Apache with mod_rewrite or PHP built-in server. No Node.js, no npm, no build step required.")
    add_empty()

    add_heading4("System Attributes")
    add_body("Portability: Runs on Windows, macOS, Linux with PHP 8.0+ and SQLite. No OS-specific dependencies.")
    add_body("Maintainability: Clean separation of concerns \u2014 config/database.php, includes/auth.php, includes/functions.php, and page files.")
    add_body("Reliability: Auto-creates database tables on first request. Graceful error handling with try-catch blocks.")
    add_body("Extensibility: New QR types require only 2 additions \u2014 one entry in getQRTypes() array and one case in encodeQRContent() switch.")
    add_empty()

    add_heading4("Hardware Requirements")
    add_body("Minimum Requirements:")
    add_bullet("Processor: Dual-core CPU (Intel i3 or equivalent)")
    add_bullet("RAM: 4 GB")
    add_bullet("Storage: 100 MB free space (SQLite database + uploads)")
    add_bullet("Internet: Required for CDN resources (Tailwind, fonts, icons, QR libraries)")
    add_bullet("Camera: Optional (required for QR scanner feature)")

    add_footer_line()
    add_page_break()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                    PRACTICAL 5                                   ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_5():
    add_uni_header()
    add_heading2("Practical 5")
    add_heading3("System Analysis and Design")
    add_body("Aim: Perform system analysis on selected system. 1) Systems analysis (what the system should do) 2) Systems design (how to accomplish the objective of the system) (Hint: Flowcharts / ER diagrams)")
    add_empty()

    add_heading3("System Analysis")
    add_body("QRCode Pro is a comprehensive QR code platform designed to generate, customize, scan, and manage QR codes. The system analysis identifies what the system should do by examining objectives, requirements, and existing gaps.")
    add_empty()

    add_heading4("System Objectives")
    add_bullet("Comprehensive QR Generation: Support 24+ QR code types with type-specific encoding (URL, WiFi, vCard, UPI, etc.).")
    add_bullet("Deep Visual Customization: Provide 62 controls across 9 categories for professional-grade QR codes.")
    add_bullet("Bulk Processing: Enable batch generation of up to 100 QR codes with CSV support and ZIP download.")
    add_bullet("Integrated Scanning: Decode QR codes via camera or image upload with content type detection.")
    add_bullet("History Management: Persistent, searchable, filterable QR code history with favorites and CSV export.")
    add_bullet("Analytics: Dashboard with statistics, type distribution charts, and activity trends.")
    add_bullet("Privacy: Self-hosted, open-source, no third-party data sharing, bcrypt password security.")
    add_empty()

    add_heading4("Functional Requirements")
    add_body("The system includes a User Authentication Module supporting registration, login (via username or email), session management, and profile editing. Passwords are hashed with bcrypt and sessions are regenerated on login to prevent fixation attacks.")
    add_body("The QR Generation Module provides 23 type-specific input forms, each with validation and real-time preview using the qr-code-styling JavaScript library. Content is encoded according to industry standards (tel: for phone, WIFI:T:... for WiFi, BEGIN:VCARD for contacts, etc.).")
    add_body("The Customization Engine offers 62 controls organized in 9 collapsible accordion sections: color presets (8 schemes), dot style & gradient (9 controls), background (8 controls), eye frame (7 controls), eye dot (7 controls), logo/watermark (5 controls), frame & label (8 controls), shadow & effects (5 controls), and size & export (5 controls).")
    add_body("The History Module provides paginated grid display (12 per page), type filtering, favorites filtering, client-side search, inline delete with confirmation, CSV export, and mini QR previews rendered on page load.")
    add_body("The Bulk Generation Module accepts manual text input or CSV upload, validates a maximum of 100 items, generates QR codes client-side with a progress bar and thumbnail previews, and packages output as a ZIP archive with a manifest file.")
    add_body("The Scanner Module uses the html5-qrcode library for camera scanning at 10fps with multi-camera support, and also supports image file upload with drag-and-drop. It auto-detects content types (URL, email, phone, WiFi, text) and maintains a session-based scan history.")
    add_empty()

    add_heading4("Non-Functional Requirements")
    add_bullet("Performance: Real-time QR preview with response time under 200ms (160ms debounce).")
    add_bullet("Scalability: Modular architecture allowing easy addition of new QR types.")
    add_bullet("Usability: Intuitive UI with animations, color-coded icons, and mobile-responsive design.")
    add_bullet("Security: Bcrypt hashing, session regeneration, input validation, directory access blocking.")
    add_bullet("Maintainability: Clean PHP includes structure, CSS custom properties, configurable base path.")
    add_bullet("Reliability: Auto-creating database schema, WAL journal mode, graceful error handling.")
    add_empty()

    add_heading4("Gaps in Existing Systems")
    add_bullet("Most free QR tools support only 3\u20135 types; QRCode Pro supports 24+.")
    add_bullet("Customization is typically limited to basic colors; QRCode Pro offers 62 controls including gradients, dot styles, eye frames, logos, shadows, and frames.")
    add_bullet("Bulk generation is rare or paid; QRCode Pro provides free bulk processing for up to 100 codes.")
    add_bullet("QR scanning is usually a separate app; QRCode Pro integrates scanning directly.")
    add_bullet("History and analytics are absent in most free tools; QRCode Pro provides full history, favorites, and Chart.js dashboards.")
    add_empty()

    add_heading3("System Design")
    add_body("System Design defines the architecture, modules, interfaces, and data flow required to meet system requirements. QRCode Pro follows a three-tier architecture designed for simplicity, performance, and extensibility.")
    add_empty()

    add_heading4("Architectural Design")
    add_body("QRCode Pro follows a three-tier architecture:")
    add_body("Presentation Layer: Tailwind CSS frontend with Lucide icons, Google Fonts (Outfit, Plus Jakarta Sans, JetBrains Mono). Responsive design with mobile bottom navigation. Dark mode support via CSS custom properties and html.dark class.")
    add_body("Business Logic Layer: PHP 8.0+ backend handling authentication (auth.php), QR content encoding (functions.php \u2014 encodeQRContent with switch-case for 23 types), CRUD operations, session management, and clean URL routing (router.php / .htaccess).")
    add_body("Data Layer: SQLite database (qrcode_pro.db) with 4 tables, WAL journal mode, foreign keys enabled. Auto-initializing schema via initDatabase() called on every request.")
    add_empty()

    add_heading4("Database Design and ER Diagram")
    add_body("The database schema consists of 4 tables with clear relationships:")
    add_empty()

    add_heading4("Entities and Relationships")
    add_body("users (id, username, email, password, full_name, avatar, plan, created_at, updated_at)", bold=True)
    add_bullet("Primary entity representing registered users.")
    add_bullet("A user can create multiple QR codes (1:N with qr_codes).")
    add_bullet("A user can create multiple bulk jobs (1:N with bulk_jobs).")
    add_empty()

    add_body("qr_codes (id, user_id, type, title, content, filename, settings, scans, is_favorite, created_at)", bold=True)
    add_bullet("Stores all generated QR codes with their encoded content and customization settings (JSON).")
    add_bullet("Foreign key: user_id references users(id) ON DELETE CASCADE.")
    add_bullet("Indexed on user_id and type for efficient querying.")
    add_empty()

    add_body("bulk_jobs (id, user_id, job_name, total_codes, zip_filename, created_at)", bold=True)
    add_bullet("Tracks bulk generation jobs and their output archives.")
    add_bullet("Foreign key: user_id references users(id) ON DELETE CASCADE.")
    add_empty()

    add_body("contact_messages (id, name, email, subject, message, is_read, created_at)", bold=True)
    add_bullet("Stores contact form submissions. Independent entity (no foreign key to users).")
    add_empty()

    # ER Diagram as text
    add_heading4("Figure 5.1: ER Diagram")
    add_empty()
    er = """
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502       users              \u2502
\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524
\u2502 PK  id (INTEGER)        \u2502
\u2502     username (TEXT, UQ)  \u2502
\u2502     email (TEXT, UQ)     \u2502
\u2502     password (TEXT)      \u2502
\u2502     full_name (TEXT)     \u2502
\u2502     plan (TEXT)          \u2502
\u2502     created_at (DATETIME)\u2502
\u2502     updated_at (DATETIME)\u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
        \u2502 1
        \u2502
        \u2502 creates (1:N)
        \u2502
        \u25bc N
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502       qr_codes                 \u2502
\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524
\u2502 PK  id (INTEGER)              \u2502
\u2502 FK  user_id (INTEGER)         \u2502
\u2502     type (TEXT)                \u2502
\u2502     title (TEXT)               \u2502
\u2502     content (TEXT)             \u2502
\u2502     filename (TEXT)            \u2502
\u2502     settings (TEXT/JSON)       \u2502
\u2502     scans (INTEGER)           \u2502
\u2502     is_favorite (INTEGER)     \u2502
\u2502     created_at (DATETIME)     \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
"""
    p = doc.add_paragraph()
    r = p.add_run(er)
    r.font.name = "Courier New"
    r.font.size = Pt(8)
    add_empty()

    # DFD
    add_heading4("Figure 5.2: Data Flow Diagram (Level 0 \u2014 Context)")
    dfd0 = """
  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510                                        \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
  \u2502  Guest   \u2502 \u2500\u2500\u2500 View Pages \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2192      \u2502           \u2502
  \u2502  User    \u2502 \u2500\u2500\u2500 Register/Login \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2192      \u2502           \u2502
  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518                          \u2502           \u2502
                                         \u2502  QRCode   \u2502 \u2500\u2500\u2500\u2192 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510                    \u2502  Pro      \u2502      \u2502  SQLite    \u2502
  \u2502 Registered\u2502 \u2500\u2500\u2500 Generate QR \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2192      \u2502  System   \u2502 \u2500\u2500\u2500\u2192 \u2502  Database  \u2502
  \u2502 User     \u2502 \u2500\u2500\u2500 Scan QR \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2192      \u2502           \u2502      \u2502  (qrcode_  \u2502
  \u2502          \u2502 \u2500\u2500\u2500 Bulk Generate \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2192      \u2502           \u2502      \u2502   pro.db)  \u2502
  \u2502          \u2502 \u2500\u2500\u2500 View History \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2192      \u2502           \u2502      \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
  \u2502          \u2502 \u2190\u2500\u2500\u2500 QR Code + Export \u2500\u2500\u2500\u2500\u2500      \u2502           \u2502
  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518                          \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
"""
    p = doc.add_paragraph()
    r = p.add_run(dfd0)
    r.font.name = "Courier New"
    r.font.size = Pt(7)
    add_empty()

    # Architecture
    add_heading4("Figure 5.3: System Architecture Diagram")
    arch = """
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502                    PRESENTATION LAYER                       \u2502
\u2502  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   \u2502
\u2502  \u2502 Tailwind  \u2502 \u2502  Lucide   \u2502 \u2502  Google   \u2502 \u2502 qr-code-  \u2502   \u2502
\u2502  \u2502 CSS (CDN) \u2502 \u2502  Icons    \u2502 \u2502  Fonts    \u2502 \u2502 styling   \u2502   \u2502
\u2502  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
                              \u2502
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502                    BUSINESS LOGIC LAYER                      \u2502
\u2502  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u2502
\u2502  \u2502 router.php  \u2502 \u2502  auth.php    \u2502 \u2502 functions.  \u2502 \u2502 api/    \u2502 \u2502
\u2502  \u2502 .htaccess   \u2502 \u2502  (session,   \u2502 \u2502 php (QR     \u2502 \u2502 generate\u2502 \u2502
\u2502  \u2502 (routing)   \u2502 \u2502  bcrypt)     \u2502 \u2502 encode,     \u2502 \u2502 .php    \u2502 \u2502
\u2502  \u2502             \u2502 \u2502             \u2502 \u2502 CRUD, CSRF) \u2502 \u2502 (REST)  \u2502 \u2502
\u2502  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
                              \u2502
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502                       DATA LAYER                             \u2502
\u2502        \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510        \u2502
\u2502        \u2502   SQLite (qrcode_pro.db)               \u2502        \u2502
\u2502        \u2502   config/database.php (PDO)             \u2502        \u2502
\u2502        \u2502   Tables: users, qr_codes,              \u2502        \u2502
\u2502        \u2502           bulk_jobs, contact_messages    \u2502        \u2502
\u2502        \u2502   WAL mode | Foreign keys ON             \u2502        \u2502
\u2502        \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518        \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
"""
    p = doc.add_paragraph()
    r = p.add_run(arch)
    r.font.name = "Courier New"
    r.font.size = Pt(7)

    add_footer_line()
    add_page_break()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                    PRACTICAL 6                                   ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_6():
    add_uni_header()
    add_heading2("Practical 6")
    add_heading3("User\u2019s View Analysis")
    add_body("Aim: Perform the user\u2019s view analysis for the suggested system by drawing Use Case Diagram.")
    add_empty()

    add_heading4("User View Analysis")
    add_body("User\u2019s view analysis is the process of understanding the interactions of different users with the system. It helps in identifying the functionalities required by different users and ensures that the system design aligns with user expectations.")
    add_empty()

    add_heading4("Actors in the System")
    add_body("An actor is an entity that interacts with the system. QRCode Pro includes the following actors:")
    add_empty()

    add_bullet("Guest User", bold_prefix="")
    add_bullet("Views the landing page, about page, privacy policy, and contact page.")
    add_bullet("Can submit the contact form.")
    add_bullet("Can register a new account or log in to an existing account.")
    add_empty()

    add_bullet("Registered User", bold_prefix="")
    add_bullet("Generates QR codes across 24+ types with type-specific forms.")
    add_bullet("Customizes QR codes using 62 controls across 9 accordion sections.")
    add_bullet("Downloads QR codes as PNG, SVG, or PDF. Prints or copies to clipboard.")
    add_bullet("Saves QR codes to history with custom titles.")
    add_bullet("Manages history: search, filter by type, toggle favorites, delete, export CSV.")
    add_bullet("Generates QR codes in bulk (up to 100) from text or CSV, downloads as ZIP.")
    add_bullet("Scans QR codes via camera (multi-camera) or image upload (drag-drop).")
    add_bullet("Views dashboard with stats, analytics charts, and quick-generate grid.")
    add_bullet("Manages profile: update name, email, change password.")
    add_empty()

    add_bullet("System (Automated)", bold_prefix="")
    add_bullet("Auto-initializes database schema on first request (CREATE IF NOT EXISTS).")
    add_bullet("Regenerates session IDs after login/registration.")
    add_bullet("Auto-detects base path for Apache subdir vs root deployment.")
    add_bullet("Renders QR code previews client-side via qr-code-styling library.")
    add_empty()

    add_heading4("Use Cases")
    add_body("For Guest Users:")
    add_bullet("UC-01: View Landing Page \u2014 Browse hero section, features, QR type grid, stats, and CTA.")
    add_bullet("UC-02: Register \u2014 Create account with username, email, password, optional full name.")
    add_bullet("UC-03: Login \u2014 Authenticate with username/email and password.")
    add_bullet("UC-04: Submit Contact Form \u2014 Send message with name, email, subject, body.")
    add_bullet("UC-05: View About/Privacy \u2014 Read team info, tech stack, and privacy policy.")
    add_empty()

    add_body("For Registered Users:")
    add_bullet("UC-06: Generate QR Code \u2014 Select type, fill form, preview in real-time.")
    add_bullet("UC-07: Customize QR Code \u2014 Apply 62 controls (colors, dots, gradients, logo, frame, shadow).")
    add_bullet("UC-08: Download QR Code \u2014 Export as PNG (with compositing), SVG, or PDF.")
    add_bullet("UC-09: Save QR to History \u2014 Store with title, type, content, settings to database.")
    add_bullet("UC-10: Print QR Code \u2014 Open print dialog with QR image and title.")
    add_bullet("UC-11: Copy QR to Clipboard \u2014 Copy QR image to system clipboard.")
    add_bullet("UC-12: View History \u2014 Browse paginated grid of saved QR codes.")
    add_bullet("UC-13: Search/Filter History \u2014 Client-side search by title/content/type; filter by type or favorites.")
    add_bullet("UC-14: Delete QR Code \u2014 Remove from history with confirmation prompt.")
    add_bullet("UC-15: Toggle Favorite \u2014 Mark/unmark QR codes as favorites.")
    add_bullet("UC-16: Export History CSV \u2014 Download filtered history as CSV file.")
    add_bullet("UC-17: Bulk Generate \u2014 Enter up to 100 items or upload CSV, generate all, download ZIP.")
    add_bullet("UC-18: Scan QR (Camera) \u2014 Start camera, detect QR code, display decoded content.")
    add_bullet("UC-19: Scan QR (Upload) \u2014 Upload/drag-drop image, decode QR content.")
    add_bullet("UC-20: View Dashboard \u2014 See stats, charts, recent codes, quick-generate grid.")
    add_bullet("UC-21: Update Profile \u2014 Change full name and email.")
    add_bullet("UC-22: Change Password \u2014 Update password with current password verification.")
    add_bullet("UC-23: Logout \u2014 Destroy session, clear cookies, redirect to login.")
    add_empty()

    add_heading4("Figure 6.1: Use Case Diagram")
    uc = """
                        \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
                        \u2502              QRCode Pro System              \u2502
                        \u2502                                              \u2502
   \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510          \u2502  (View Landing)  (View About/Privacy)       \u2502
   \u2502 Guest \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2502  (Register)       (Login)                   \u2502
   \u2502 User  \u2502          \u2502  (Submit Contact Form)                      \u2502
   \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518          \u2502                                              \u2502
                        \u2502  (Generate QR)    (Customize QR - 62 ctrl)  \u2502
   \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510       \u2502  (Download PNG/SVG/PDF)  (Print QR)          \u2502
   \u2502Registered\u2502\u2500\u2500\u2500\u2500\u2500\u2502  (Copy to Clipboard)    (Save to History)    \u2502
   \u2502  User    \u2502       \u2502  (View History)  (Search/Filter History)   \u2502
   \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518       \u2502  (Delete QR)  (Toggle Favorite)  (CSV)    \u2502
                        \u2502  (Bulk Generate)   (Download ZIP)           \u2502
                        \u2502  (Scan Camera)     (Scan Upload)            \u2502
                        \u2502  (View Dashboard)  (View Analytics)         \u2502
                        \u2502  (Update Profile)  (Change Password)        \u2502
   \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510         \u2502  (Logout)                                   \u2502
   \u2502 System \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2502                                              \u2502
   \u2502 (Auto) \u2502         \u2502  (Auto-init DB)  (Regenerate Session)     \u2502
   \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518         \u2502  (Auto-detect Base Path) (Render QR)       \u2502
                        \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
"""
    p = doc.add_paragraph()
    r = p.add_run(uc)
    r.font.name = "Courier New"
    r.font.size = Pt(7)

    add_footer_line()
    add_page_break()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                    PRACTICAL 7                                   ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_7():
    add_uni_header()
    add_heading2("Practical 7")
    add_heading3("Structural View Diagrams")
    add_body("Aim: Design structural view diagram for the selected system using Class Diagram, Object Diagram, and Component Diagram.")
    add_empty()

    add_heading4("Class Diagram")
    add_body("The Class Diagram shows the static structure of the QRCode Pro system, depicting classes, their attributes, operations, and relationships.")
    add_empty()

    add_heading4("Figure 7.1: Class Diagram")
    cd = """
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510    \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502            User                 \u2502    \u2502             QRCode                 \u2502
\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524    \u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524
\u2502 - id: int                      \u2502    \u2502 - id: int                         \u2502
\u2502 - username: string (UQ)        \u2502    \u2502 - user_id: int (FK)               \u2502
\u2502 - email: string (UQ)           \u2502    \u2502 - type: string                    \u2502
\u2502 - password: string (bcrypt)    \u2502    \u2502 - title: string                   \u2502
\u2502 - full_name: string            \u2502    \u2502 - content: string                 \u2502
\u2502 - plan: string                 \u2502    \u2502 - filename: string                \u2502
\u2502 - created_at: datetime         \u2502    \u2502 - settings: JSON                  \u2502
\u2502 - updated_at: datetime         \u2502    \u2502 - scans: int                      \u2502
\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524    \u2502 - is_favorite: boolean           \u2502
\u2502 + register(user, email, pass)  \u2502    \u2502 - created_at: datetime            \u2502
\u2502 + login(login, password)       \u2502    \u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524
\u2502 + logout()                     \u2502    \u2502 + save(userId, type, content)      \u2502
\u2502 + getCurrentUser(): array      \u2502    \u2502 + getUserQRCodes(userId): array[]  \u2502
\u2502 + getUserStats(id): array      \u2502    \u2502 + delete(id, userId): bool         \u2502
\u2502 + updateProfile(name, email)   \u2502    \u2502 + toggleFavorite(id, userId): bool \u2502
\u2502 + changePassword(old, new)     \u2502    \u2502 + encodeContent(type, data): str   \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518    \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
        \u2502 1:N                               \u2502 (composition)
        \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518

\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510    \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502          BulkJob              \u2502    \u2502         ContactMessage             \u2502
\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524    \u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524
\u2502 - id: int                    \u2502    \u2502 - id: int                         \u2502
\u2502 - user_id: int (FK)          \u2502    \u2502 - name: string                    \u2502
\u2502 - job_name: string           \u2502    \u2502 - email: string                   \u2502
\u2502 - total_codes: int           \u2502    \u2502 - subject: string                 \u2502
\u2502 - zip_filename: string       \u2502    \u2502 - message: string                 \u2502
\u2502 - created_at: datetime       \u2502    \u2502 - is_read: boolean                \u2502
\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524    \u2502 - created_at: datetime            \u2502
\u2502 + create(userId, items)      \u2502    \u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524
\u2502 + downloadZip(): blob        \u2502    \u2502 + submit(name, email, msg)         \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518    \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
"""
    p = doc.add_paragraph()
    r = p.add_run(cd)
    r.font.name = "Courier New"
    r.font.size = Pt(7)
    add_empty()

    # Object Diagram
    add_heading4("Object Diagram")
    add_body("An Object Diagram showing a specific instance of a user generating a WiFi QR code:")
    add_empty()
    add_heading4("Figure 7.2: Object Diagram")
    od = """
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502   aadil : User                \u2502     \u2502   wifiQR : QRCode                    \u2502
\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524     \u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524
\u2502 id = 1                       \u2502     \u2502 id = 42                              \u2502
\u2502 username = "aadilp"          \u2502\u2500\u2500\u2500\u2500\u25b6\u2502 user_id = 1                          \u2502
\u2502 email = "aadil@mu.edu"       \u2502     \u2502 type = "wifi"                        \u2502
\u2502 full_name = "Aadil Parmar"   \u2502     \u2502 title = "Office WiFi"                \u2502
\u2502 plan = "free"                \u2502     \u2502 content = "WIFI:T:WPA;S:OffNet;P:..."  \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518     \u2502 settings = {"fg":"#4F46E5","bg":"#EEF"} \u2502
                                      \u2502 is_favorite = true                   \u2502
                                      \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
"""
    p = doc.add_paragraph()
    r = p.add_run(od)
    r.font.name = "Courier New"
    r.font.size = Pt(7)
    add_empty()

    # Component Diagram
    add_heading4("Component Diagram")
    add_body("The Component Diagram shows the physical components of the QRCode Pro system and their dependencies:")
    add_empty()
    add_heading4("Figure 7.3: Component Diagram")
    comp = """
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  \u00abFrontend Components\u00bb                                       \u2502
\u2502  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u2502
\u2502  \u2502 header.php \u2502 \u2502 footer.php \u2502 \u2502 generate.  \u2502 \u2502 scanner. \u2502 \u2502
\u2502  \u2502 (nav,dark \u2502 \u2502 (links,   \u2502 \u2502 php (62   \u2502 \u2502 php      \u2502 \u2502
\u2502  \u2502  mode,CSS)\u2502 \u2502  icons)   \u2502 \u2502 controls) \u2502 \u2502 (camera) \u2502 \u2502
\u2502  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
                              \u2502 depends on
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  \u00abBackend Components\u00bb                                        \u2502
\u2502  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u2502
\u2502  \u2502 router.php \u2502 \u2502  auth.php  \u2502 \u2502 functions. \u2502 \u2502 api/     \u2502 \u2502
\u2502  \u2502 .htaccess  \u2502 \u2502 (session, \u2502 \u2502 php (QR   \u2502 \u2502 generate \u2502 \u2502
\u2502  \u2502 (routing)  \u2502 \u2502  bcrypt)  \u2502 \u2502 encode)   \u2502 \u2502 .php     \u2502 \u2502
\u2502  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
                              \u2502 connects to
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  \u00abData Store\u00bb                                                 \u2502
\u2502  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   \u2502
\u2502  \u2502 config/        \u2502  \u2502 SQLite: data/qrcode_pro.db       \u2502   \u2502
\u2502  \u2502 database.php   \u2502  \u2502 Tables: users, qr_codes,         \u2502   \u2502
\u2502  \u2502 (PDO, schema)  \u2502  \u2502         bulk_jobs,contact_msgs    \u2502   \u2502
\u2502  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
"""
    p = doc.add_paragraph()
    r = p.add_run(comp)
    r.font.name = "Courier New"
    r.font.size = Pt(7)

    add_footer_line()
    add_page_break()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                    PRACTICAL 8                                   ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_8():
    add_uni_header()
    add_heading2("Practical 8")
    add_heading3("Function-Oriented Diagrams (DFD)")
    add_body("Aim: Design function-oriented diagram for the selected system using Data Flow Diagrams.")
    add_empty()

    add_heading4("Level 0 DFD (Context Diagram)")
    add_body("The Level 0 DFD shows QRCode Pro as a single process interacting with external entities:")
    add_empty()

    add_heading4("Figure 8.1: Level 0 DFD")
    add_body("[Guest User] \u2500\u2500 Registration Data / Contact Form \u2500\u2500\u25b6 [QRCode Pro System] \u2500\u25b6 [SQLite Database]")
    add_body("[Registered User] \u2500\u2500 QR Data, Customizations, Bulk Input, Scan Image \u2500\u2500\u25b6 [QRCode Pro System]")
    add_body("[QRCode Pro System] \u2500\u2500 QR Code (PNG/SVG/PDF), History, Dashboard, Scan Result \u2500\u2500\u25b6 [User]")
    add_empty()

    add_heading4("Level 1 DFD")
    add_body("The Level 1 DFD decomposes QRCode Pro into its major subsystems:")
    add_empty()

    add_heading4("Figure 8.2: Level 1 DFD")
    dfd1 = """
                     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
                     \u2502   1.0 Auth       \u2502
  [User] \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502   Module         \u2502\u2500\u2500\u2500\u25b6 ||users||
   login/register    \u2502   (auth.php)     \u2502
                     \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
                            \u2502 session
                            \u25bc
                     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
  [User] \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502   2.0 QR         \u2502
   QR data + 62      \u2502   Generation     \u2502\u2500\u2500\u2500\u25b6 ||qr_codes||
   customization      \u2502   (generate.php) \u2502
   controls          \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
                            \u2502 QR image
                            \u25bc
                     [User] \u2190 PNG/SVG/PDF download

                     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
  [User] \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502   3.0 History    \u2502
   search/filter      \u2502   Module        \u2502\u2500\u2500\u2500\u25b6 ||qr_codes||
                     \u2502   (history.php)  \u2502\u2500\u2500\u2500\u25b6 [User] CSV
                     \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518

                     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
  [User] \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502   4.0 Bulk       \u2502
   text list / CSV    \u2502   Generator     \u2502\u2500\u2500\u2500\u25b6 [User] ZIP
                     \u2502   (bulk.php)     \u2502
                     \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518

                     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
  [User] \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502   5.0 Scanner    \u2502
   camera / image     \u2502   Module        \u2502\u2500\u2500\u2500\u25b6 [User] decoded text
                     \u2502   (scanner.php)  \u2502
                     \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518

                     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
  [User] \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502   6.0 Dashboard  \u2502\u2500\u2500\u2500\u25b6 ||qr_codes||
                     \u2502   (dashboard.php)\u2502\u2500\u2500\u2500\u25b6 [User] stats+charts
                     \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
"""
    p = doc.add_paragraph()
    r = p.add_run(dfd1)
    r.font.name = "Courier New"
    r.font.size = Pt(7)

    add_footer_line()
    add_page_break()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                    PRACTICAL 9                                   ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_9():
    add_uni_header()
    add_heading2("Practical 9")
    add_heading3("Behavioral View Diagrams")
    add_body("Aim: Design the behavioral view diagram for the selected system using Interaction Diagrams (Sequence, Collaboration), State-Chart Diagrams, and Activity Diagrams.")
    add_empty()

    # Sequence Diagram
    add_heading4("Sequence Diagram")
    add_body("The Sequence Diagram shows the time-ordered interaction between the user and QRCode Pro system components during QR code generation:")
    add_empty()

    add_heading4("Figure 9.1: Sequence Diagram \u2014 QR Code Generation Flow")
    seq = """
  User          Browser/JS       generate.php     functions.php    auth.php       SQLite DB
   \u2502               \u2502                 \u2502                \u2502               \u2502               \u2502
   \u2502\u2500\u2500open /generate\u25b6\u2502                 \u2502                \u2502               \u2502               \u2502
   \u2502               \u2502\u2500\u2500GET request\u2500\u2500\u2500\u2500\u25b6\u2502                \u2502               \u2502               \u2502
   \u2502               \u2502                 \u2502\u2500\u2500requireLogin()\u2500\u25b6\u2502               \u2502               \u2502
   \u2502               \u2502                 \u2502                \u2502\u2500\u2500check session\u2500\u25b6\u2502               \u2502
   \u2502               \u2502                 \u2502                \u2502               \u2502\u2500\u2500SELECT user\u2500\u25b6\u2502
   \u2502               \u2502                 \u2502                \u2502               \u2502\u2190\u2500\u2500user data\u2500\u2500\u2524
   \u2502               \u2502                 \u2502\u2190\u2500\u2500getQRTypes()\u2500\u2500\u2524               \u2502               \u2502
   \u2502               \u2502\u2190\u2500\u2500HTML + 62 controls\u2524                \u2502               \u2502               \u2502
   \u2502               \u2502                 \u2502                \u2502               \u2502               \u2502
   \u2502\u2500\u2500fill form\u2500\u2500\u2500\u25b6\u2502                 \u2502                \u2502               \u2502               \u2502
   \u2502               \u2502\u2500\u2500updateQR()\u2500\u2500\u2500\u25b6\u2502(client-side)   \u2502               \u2502               \u2502
   \u2502               \u2502  debounce 160ms \u2502                \u2502               \u2502               \u2502
   \u2502               \u2502\u2500\u2500getQRContent()\u25b6\u2502                \u2502               \u2502               \u2502
   \u2502               \u2502\u2500\u2500getQROptions()\u25b6\u2502(reads 62 ctrl) \u2502               \u2502               \u2502
   \u2502               \u2502\u2500\u2500QRCodeStyling()\u25b6 canvas render  \u2502               \u2502               \u2502
   \u2502\u2190\u2500\u2500live preview\u2500\u2524                 \u2502                \u2502               \u2502               \u2502
   \u2502               \u2502                 \u2502                \u2502               \u2502               \u2502
   \u2502\u2500\u2500click Save\u2500\u2500\u25b6\u2502                 \u2502                \u2502               \u2502               \u2502
   \u2502               \u2502\u2500\u2500POST form\u2500\u2500\u2500\u2500\u2500\u25b6\u2502                \u2502               \u2502               \u2502
   \u2502               \u2502                 \u2502\u2500\u2500saveQRCode()\u2500\u2500\u25b6\u2502               \u2502               \u2502
   \u2502               \u2502                 \u2502                \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502\u2500\u2500INSERT qr\u2500\u2500\u25b6\u2502
   \u2502               \u2502                 \u2502                \u2502               \u2502\u2190\u2500\u2500success\u2500\u2500\u2500\u2524
   \u2502               \u2502                 \u2502\u2190\u2500\u2500setFlash(ok)\u2500\u2500\u2524               \u2502               \u2502
   \u2502               \u2502\u2190\u2500\u2500redirect /history\u2524                \u2502               \u2502               \u2502
   \u2502\u2190\u2500\u2500history page\u2500\u2524                 \u2502                \u2502               \u2502               \u2502
"""
    p = doc.add_paragraph()
    r = p.add_run(seq)
    r.font.name = "Courier New"
    r.font.size = Pt(6)
    add_empty()

    # Collaboration Diagram
    add_heading4("Figure 9.2: Collaboration Diagram \u2014 QR Generation")
    add_body("1: User \u2500\u25b6 Browser: fill QR form")
    add_body("2: Browser \u2500\u25b6 JS Engine: updateQR() [debounced]")
    add_body("3: JS Engine \u2500\u25b6 QRCodeStyling: render(options)")
    add_body("4: QRCodeStyling \u2500\u25b6 Canvas: draw QR code")
    add_body("5: User \u2500\u25b6 Browser: click Save")
    add_body("6: Browser \u2500\u25b6 generate.php: POST form data")
    add_body("7: generate.php \u2500\u25b6 functions.php: saveQRCode()")
    add_body("8: functions.php \u2500\u25b6 database.php: INSERT INTO qr_codes")
    add_body("9: generate.php \u2500\u25b6 Browser: redirect to /history")
    add_empty()

    # State Chart
    add_heading4("Figure 9.3: State-Chart Diagram \u2014 QR Code Lifecycle")
    sc = """
  [\u25cf] \u2500\u25b6 [IDLE: No Content]
              \u2502
              \u2502 user types input
              \u25bc
        [PREVIEWING: Live QR]
              \u2502
         \u250c\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
         \u2502   \u2502       \u2502
     customize  \u2502   change type
     (62 ctrl)  \u2502   (select pill)
         \u2502   \u2502       \u2502
         \u25bc   \u2502       \u25bc
  [CUSTOMIZED]  \u2502  [TYPE CHANGED]
         \u2502   \u2502       \u2502
         \u2514\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
              \u2502
     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
     \u2502      \u2502          \u2502
  Download   Save     Print/Copy
     \u2502      \u2502          \u2502
     \u25bc      \u25bc          \u25bc
[EXPORTED] [SAVED]  [PRINTED/COPIED]
  PNG/SVG   to DB    to printer/
  /PDF     history   clipboard
"""
    p = doc.add_paragraph()
    r = p.add_run(sc)
    r.font.name = "Courier New"
    r.font.size = Pt(8)
    add_empty()

    # Activity Diagram
    add_heading4("Figure 9.4: Activity Diagram \u2014 Complete User Workflow")
    act = """
  [\u25cf START]
      \u2502
      \u25bc
  [Visit Landing Page]
      \u2502
      \u25c6 Logged in?
     / \\
   No   Yes
    \u2502     \u2502
    \u25bc     \u25bc
 [Register/Login]   [Dashboard]
    \u2502                 \u2502
    \u25bc            \u250c\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
 [Dashboard]    \u2502    \u2502             \u2502
                \u25bc    \u25bc             \u25bc
          [Generate] [Bulk]     [Scanner]
             \u2502        \u2502            \u2502
             \u25bc        \u25bc            \u25bc
      [Select QR Type] [Enter Items] [Start Camera]
             \u2502        \u2502            \u2502
             \u25bc        \u25bc            \u25bc
      [Fill Form]   [Generate All] [Decode QR]
             \u2502        \u2502            \u2502
             \u25bc        \u25bc            \u25bc
      [Customize 62] [Download ZIP] [Show Result]
             \u2502
        \u250c\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
        \u2502    \u2502        \u2502
        \u25bc    \u25bc        \u25bc
    [Download] [Save]  [Print]
    PNG/SVG/PDF to DB  /Copy
        \u2502    \u2502        \u2502
        \u2514\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
             \u2502
             \u25bc
         [History] \u2500\u25b6 Search/Filter/Favorite/Delete/CSV
             \u2502
             \u25bc
         [\u25cf END]
"""
    p = doc.add_paragraph()
    r = p.add_run(act)
    r.font.name = "Courier New"
    r.font.size = Pt(8)

    add_footer_line()
    add_page_break()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                  PRACTICAL 10                                    ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_10():
    add_uni_header()
    add_heading2("Practical 10")
    add_body("Development and Management of Software using Reverse Engineering", bold=True)
    add_empty()

    add_heading4("Step 1: Existing Solution and URL")
    add_body("Software: QR Monkey (Free QR Code Generator)")
    add_body("URL: https://www.qrmonkey.com")
    add_empty()

    add_heading4("Step 2: Existing Features with Details")
    add_bullet("Generate QR codes for URL, text, email, phone, SMS, WiFi, vCard")
    add_bullet("Basic color customization (foreground + background)")
    add_bullet("Logo upload in center of QR code")
    add_bullet("Download as PNG or SVG")
    add_bullet("No user accounts required")
    add_bullet("Mobile-responsive web interface")
    add_empty()

    add_heading4("Step 3: Technology Used for Existing Solution")
    add_bullet("Web-based SPA (JavaScript)")
    add_bullet("Client-side QR generation (Canvas API)")
    add_bullet("No backend database (stateless)")
    add_bullet("CDN-hosted static assets")
    add_bullet("No server-side processing for QR generation")
    add_empty()

    add_heading4("Step 4: Reasons to Perform Reverse Engineering")
    add_bullet("Understand how client-side QR generation works with canvas APIs")
    add_bullet("Identify limitations in customization depth (only basic colors)")
    add_bullet("Analyze the lack of user history, analytics, and bulk generation")
    add_bullet("Study the stateless architecture and its privacy implications")
    add_bullet("Build a more comprehensive solution based on identified gaps")
    add_empty()

    add_heading4("Step 5: Drawbacks of Existing Solution")
    add_bullet("Limited to ~8 QR types (no UPI, crypto, social media, calendar, MeCard)")
    add_bullet("No dot style, eye frame, gradient, or shadow customization")
    add_bullet("No bulk generation capability")
    add_bullet("No QR scanner integration")
    add_bullet("No user accounts, history, or favorites")
    add_bullet("No dashboard or analytics")
    add_bullet("No PDF export or print functionality")
    add_empty()

    add_heading4("Step 6: Features Added / Updated / Replaced / Removed")
    add_body("Added:", bold=True)
    add_bullet("24+ QR types (UPI, Bitcoin, social media, calendar, GPS, MeCard, App Store)")
    add_bullet("62 customization controls (dot styles, gradients, eye frames, logos, shadows, frames)")
    add_bullet("Bulk generation (100 codes, CSV, ZIP download)")
    add_bullet("QR scanner (camera + upload)")
    add_bullet("User accounts with history, favorites, and dashboard analytics")
    add_body("Updated:", bold=True)
    add_bullet("Color system \u2192 Full gradient support with angle control")
    add_bullet("Logo upload \u2192 Logo with size, margin, opacity, and hide-dots controls")
    add_body("Replaced:", bold=True)
    add_bullet("Stateless architecture \u2192 PHP + SQLite with persistent user data")
    add_bullet("Basic Canvas QR \u2192 qr-code-styling library with advanced rendering")
    add_body("Removed:", bold=True)
    add_bullet("Ads and tracking (QRCode Pro is ad-free and self-hosted)")
    add_empty()

    add_heading4("Step 7: Technologies Used")
    add_bullet("Backend: PHP 8.0+")
    add_bullet("Database: SQLite (zero-config)")
    add_bullet("Frontend: Tailwind CSS (CDN), Vanilla JavaScript")
    add_bullet("QR Engine: qr-code-styling (client-side)")
    add_bullet("Charts: Chart.js")
    add_bullet("Scanner: html5-qrcode")
    add_bullet("Export: jsPDF, JSZip, FileSaver.js")
    add_empty()

    add_heading4("Step 8: Advantages of Updated Solution")
    add_bullet("24+ QR types vs ~8 in QR Monkey")
    add_bullet("62 customization controls vs basic colors only")
    add_bullet("Persistent history with search, filter, favorites")
    add_bullet("Bulk generation with ZIP download")
    add_bullet("Integrated QR scanner")
    add_bullet("Self-hosted, open-source, privacy-respecting")
    add_empty()

    add_heading4("Step 9: Blueprint of Updated Solution")
    add_bullet("Pages: Landing, Login, Register, Dashboard, Generate, History, Bulk, Scanner, Profile, About, Contact, Privacy")
    add_bullet("Flow: Register \u2192 Login \u2192 Dashboard \u2192 Generate (select type \u2192 fill form \u2192 customize \u2192 preview \u2192 download/save)")
    add_empty()

    add_heading4("Step 10: Comparison")
    cmp_headers = ["Feature", "QR Monkey (Existing)", "QRCode Pro (New)"]
    cmp_rows = [
        ["QR Types", "~8 types", "24+ types"],
        ["Customization", "Basic colors + logo", "62 controls, 9 sections"],
        ["Bulk Generation", "Not available", "Up to 100, CSV + ZIP"],
        ["QR Scanner", "Not available", "Camera + upload"],
        ["User Accounts", "None", "Full auth system"],
        ["History", "None", "Paginated, searchable"],
        ["Analytics", "None", "Chart.js dashboard"],
        ["Export Formats", "PNG, SVG", "PNG, SVG, PDF, Print, Copy"],
        ["Dark Mode", "No", "Yes (full support)"],
        ["Self-Hosted", "No (SaaS)", "Yes (open-source)"],
    ]
    make_table(cmp_headers, cmp_rows, [1.5, 2.2, 2.5])

    add_footer_line()
    add_page_break()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                  PRACTICAL 11                                    ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_11():
    add_uni_header()
    add_heading2("Practical 11")
    add_body("Development and Management of Software using Forward Engineering", bold=True)
    add_empty()

    add_heading4("Step 1: Existing Solution and URL")
    add_body("Software: GoQR.me (QR Code API & Generator)")
    add_body("URL: https://goqr.me")
    add_empty()

    add_heading4("Step 2: Existing Features with Details")
    add_bullet("Basic QR code generation for URL, text, phone, SMS, vCard")
    add_bullet("Simple color customization")
    add_bullet("REST API for QR generation (developer-focused)")
    add_bullet("Download as PNG, SVG, EPS")
    add_bullet("No user accounts")
    add_bullet("Minimal UI with limited interactivity")
    add_empty()

    add_heading4("Step 3: Technology Used for Existing Solution")
    add_bullet("Server-side QR generation (Java backend)")
    add_bullet("REST API architecture")
    add_bullet("Static HTML/CSS frontend")
    add_bullet("No client-side rendering")
    add_bullet("No database for user data")
    add_empty()

    add_heading4("Step 4: Drawbacks of Existing Solution")
    add_bullet("Very limited QR types (no UPI, crypto, social media)")
    add_bullet("No visual customization (no dot styles, gradients, logos)")
    add_bullet("No user accounts or history management")
    add_bullet("No bulk generation capability")
    add_bullet("No QR scanner")
    add_bullet("Outdated UI design")
    add_empty()

    add_heading4("Step 5: Forward Engineering Methodology")
    add_body("Methodology: Agile-Waterfall Hybrid", bold=True)
    add_bullet("Requirements Analysis: Studied GoQR.me and 5+ other tools to define comprehensive feature set.")
    add_bullet("System Design: Designed 3-tier architecture with PHP backend, SQLite storage, and JS frontend.")
    add_bullet("Incremental Development: Built core (auth + generate) first, then added bulk, scanner, dashboard.")
    add_bullet("Testing: Cross-browser, mobile responsiveness, dark mode, 23 QR type encoding tests.")
    add_bullet("Deployment: Zero-config deployment on any PHP 8.0+ server with SQLite.")
    add_bullet("Maintenance: Git-based version control, modular architecture for easy updates.")
    add_empty()

    add_heading4("Step 6: Features Added")
    add_bullet("24+ QR types (vs 5 in GoQR.me)")
    add_bullet("62 visual customization controls")
    add_bullet("User authentication with bcrypt and session security")
    add_bullet("Persistent history with search, filter, favorites, CSV export")
    add_bullet("Bulk generation (100 codes, ZIP download)")
    add_bullet("Integrated QR scanner (camera + upload)")
    add_bullet("Dashboard with Chart.js analytics")
    add_bullet("Dark mode with localStorage persistence")
    add_bullet("PDF export, print, and clipboard copy")
    add_empty()

    add_heading4("Step 7: Technologies Used")
    add_bullet("Frontend: Tailwind CSS (CDN), JavaScript (qr-code-styling, Chart.js, html5-qrcode)")
    add_bullet("Backend: PHP 8.0+ with PDO")
    add_bullet("Database: SQLite (zero-config, WAL mode)")
    add_bullet("Export: jsPDF, JSZip, FileSaver.js")
    add_bullet("Security: bcrypt, session regeneration, .htaccess directory blocking")
    add_empty()

    add_heading4("Step 8: Advantages of Updated Solution")
    add_bullet("Comprehensive feature set in a single, unified application")
    add_bullet("Self-hosted and open-source (complete data ownership)")
    add_bullet("Zero-configuration deployment (no npm, no build step)")
    add_bullet("Modern, responsive UI with dark mode")
    add_bullet("Extensible architecture (add new QR types in 2 code locations)")
    add_empty()

    add_heading4("Step 9: Blueprint of New Solution")
    add_body("Main Structure:")
    add_bullet("Auth: Login, Register (public)")
    add_bullet("Dashboard: Stats, charts, quick-generate, recent codes")
    add_bullet("Generate: 23 type forms + 62 controls + live preview")
    add_bullet("History: Paginated grid, search, filter, favorites, CSV export")
    add_bullet("Bulk: Manual/CSV input, progress bar, ZIP download")
    add_bullet("Scanner: Camera/upload tabs, result display, scan history")
    add_bullet("Profile: Edit name/email, change password")
    add_body("Layout:", bold=True)
    add_bullet("Top navbar with logo, navigation, dark mode toggle, user menu")
    add_bullet("Mobile bottom navigation bar")
    add_bullet("Content area with responsive grid layouts")
    add_body("Flow:", bold=True)
    add_body("Register \u2192 Login \u2192 Dashboard \u2192 Generate/Bulk/Scan \u2192 Download/Save \u2192 History")
    add_empty()

    add_heading4("Step 10: Comparison")
    cmp_headers = ["Feature", "GoQR.me (Existing)", "QRCode Pro (New)"]
    cmp_rows = [
        ["QR Types", "5 types", "24+ types"],
        ["Customization", "Basic colors", "62 controls, gradients, logos"],
        ["Architecture", "Server-side Java", "Client-side JS + PHP backend"],
        ["User Accounts", "None", "Full auth with sessions"],
        ["History", "None", "Persistent, searchable, favorites"],
        ["Bulk Generation", "None", "100 codes, CSV, ZIP"],
        ["Scanner", "None", "Camera + upload"],
        ["Export", "PNG, SVG, EPS", "PNG, SVG, PDF, Print, Copy"],
        ["Dark Mode", "No", "Full support"],
        ["Deployment", "Cloud SaaS", "Self-hosted, zero-config"],
    ]
    make_table(cmp_headers, cmp_rows, [1.5, 2.2, 2.5])

    add_footer_line()
    add_page_break()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                  PRACTICAL 12                                    ║
# ╚══════════════════════════════════════════════════════════════════╝
def practical_12():
    add_uni_header()
    add_heading2("Practical 12")
    add_body("Development and Management of Software using Re-Engineering", bold=True)
    add_empty()

    add_heading4("Step 1: Existing Solution and URL")
    add_body("Software: QRCode.js (Basic JavaScript QR Library)")
    add_body("URL: https://davidshimjs.github.io/qrcodejs/")
    add_empty()

    add_heading4("Step 2: Existing Features with Details")
    add_bullet("Generate QR codes from text/URL input")
    add_bullet("Basic size configuration")
    add_bullet("Foreground and background color options")
    add_bullet("Canvas and SVG rendering modes")
    add_bullet("Lightweight JavaScript library (~15KB)")
    add_empty()

    add_heading4("Step 3: Technology Used for Existing Solution")
    add_bullet("Pure JavaScript library (no framework)")
    add_bullet("HTML5 Canvas API")
    add_bullet("No backend or database")
    add_bullet("Client-side only")
    add_empty()

    add_heading4("Step 4: Reasons to Perform Re-Engineering")
    add_bullet("Core QR rendering is functional but lacks modern features")
    add_bullet("No dot style or eye frame customization")
    add_bullet("No gradient, logo, shadow, or frame support")
    add_bullet("No application wrapper (just a library)")
    add_bullet("Need to preserve the fast, client-side rendering approach")
    add_empty()

    add_heading4("Step 5: Drawbacks of Existing Solution")
    add_bullet("Only generates basic square QR codes")
    add_bullet("No dot style variety (squares only)")
    add_bullet("No gradient or multi-color support")
    add_bullet("No logo/watermark overlay capability")
    add_bullet("No frame, label, or shadow effects")
    add_bullet("No application UI \u2014 just a JavaScript library")
    add_empty()

    add_heading4("Step 6: Things Preserved from Existing Solution")
    add_bullet("Client-side rendering approach (fast, no server round-trip)")
    add_bullet("Canvas-based output (for PNG export)")
    add_bullet("Lightweight architecture")
    add_bullet("Simple integration via CDN")
    add_body("(These core principles are preserved in QRCode Pro)")
    add_empty()

    add_heading4("Step 7: Features Added / Updated / Replaced / Removed")
    add_body("Added:", bold=True)
    add_bullet("62 visual customization controls across 9 categories")
    add_bullet("6 dot styles (square, dots, rounded, extra-rounded, classy, classy-rounded)")
    add_bullet("Gradient support for dots, background, eye frames, and eye dots")
    add_bullet("Logo overlay with size, margin, opacity, and hide-dots controls")
    add_bullet("Frame & label system with color, width, padding, radius")
    add_bullet("Shadow effects with color, blur, and offset controls")
    add_bullet("Full web application with auth, history, dashboard, bulk, scanner")
    add_body("Updated:", bold=True)
    add_bullet("Basic QR rendering \u2192 qr-code-styling library (advanced canvas)")
    add_bullet("Simple text input \u2192 23 type-specific forms with validation")
    add_body("Replaced:", bold=True)
    add_bullet("QRCode.js (basic) \u2192 qr-code-styling (advanced) for generate page")
    add_bullet("No backend \u2192 PHP 8.0+ with SQLite for persistence")
    add_body("Removed:", bold=True)
    add_bullet("Nothing fundamental \u2014 QRCode.js is still used for simple previews (history, landing page)")
    add_empty()

    add_heading4("Step 8: Technologies Used")
    add_bullet("Frontend: Tailwind CSS (CDN), qr-code-styling, QRCode.js, Chart.js, html5-qrcode")
    add_bullet("Backend: PHP 8.0+ with PDO")
    add_bullet("Database: SQLite with WAL mode")
    add_bullet("Export: jsPDF (PDF), JSZip (bulk), FileSaver.js")
    add_empty()

    add_heading4("Step 9: Advantages of Updated Solution")
    add_bullet("Full application vs just a library")
    add_bullet("62 customization controls for professional QR codes")
    add_bullet("Persistent user data with history and analytics")
    add_bullet("Bulk generation and scanning capabilities")
    add_bullet("Modern, responsive UI with dark mode")
    add_bullet("Still maintains fast, client-side rendering approach")
    add_empty()

    add_heading4("Step 10: Blueprint of Updated Solution")
    add_body("Main Layout:")
    add_bullet("Top navbar (logo, nav links, dark mode, user menu)")
    add_bullet("Mobile bottom navigation")
    add_bullet("Content area with card-based UI")
    add_bullet("Footer with links and credits")
    add_body("Flow:", bold=True)
    add_body("Open App \u2192 Register/Login \u2192 Dashboard \u2192 Generate (select type \u2192 fill form \u2192 customize 62 controls \u2192 live preview \u2192 download/save) \u2192 History")
    add_empty()

    add_heading4("Step 11: Comparison")
    cmp_headers = ["Feature", "QRCode.js (Existing)", "QRCode Pro (Re-Engineered)"]
    cmp_rows = [
        ["Type", "JavaScript library", "Full web application"],
        ["QR Types", "Text/URL only", "24+ types"],
        ["Customization", "Size + 2 colors", "62 controls, 9 sections"],
        ["Dot Styles", "Square only", "6 styles + gradients"],
        ["Logo Support", "No", "Yes (size, margin, opacity)"],
        ["Frame/Shadow", "No", "Yes (full control)"],
        ["User Accounts", "N/A", "Full auth system"],
        ["History", "N/A", "Persistent, searchable"],
        ["Bulk Generation", "N/A", "100 codes, ZIP"],
        ["Scanner", "N/A", "Camera + upload"],
        ["Dark Mode", "N/A", "Full support"],
    ]
    make_table(cmp_headers, cmp_rows, [1.5, 2.0, 2.5])

    add_footer_line()


# ══════════════════════════════════════════════════════════════════
#  MAIN — Generate the document
# ══════════════════════════════════════════════════════════════════
print("Generating lab manual...")

practical_1()
practical_2()
practical_3()
practical_4()
practical_5()
practical_6()
practical_7()
practical_8()
practical_9()
practical_10()
practical_11()
practical_12()

doc.save(OUT_FILE)
print(f"\nDone! Lab manual saved to:\n  {OUT_FILE}")
print(f"  Size: {os.path.getsize(OUT_FILE) / 1024:.1f} KB")
