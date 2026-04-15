#!/usr/bin/env python3
"""Fix Practical 1 only — no overflow, clean content, proper footer."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt

DOC_PATH = r'F:\mr_qr_v1\92200103068_EC3_Lab Manual.docx'
doc = Document(DOC_PATH)

def set_text(para, new_text):
    """Replace paragraph text keeping first run's formatting."""
    if not para.runs:
        r = para.add_run(new_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""

def set_cell_text(cell, new_text):
    if cell.paragraphs:
        set_text(cell.paragraphs[0], new_text)
        for p in cell.paragraphs[1:]:
            set_text(p, "")

# ═══════════════════════════════════════════════════════════════
# FIX 1: Shorten paragraphs BEFORE footer (P12, P16, P20, P26-28)
#         so content fits on page 1 without overflow
# ═══════════════════════════════════════════════════════════════

# P12 — Introduction (was too long, shorten)
set_text(doc.paragraphs[12],
    "QRCode Pro is a web-based QR code generation platform built with PHP 8.0+, SQLite, "
    "Tailwind CSS, and JavaScript. It supports 24+ QR code types, 62 customization controls, "
    "bulk generation (up to 100 codes), an integrated camera scanner, and full export in PNG, SVG, and PDF.")

# P16 — Problem Statement (shorten)
set_text(doc.paragraphs[16],
    "QR codes are ubiquitous in payments, marketing, and information sharing. However, existing free tools "
    "support only 3\u20135 types, offer minimal customization, lack bulk generation, provide no history or analytics, "
    "and require separate apps for scanning. There is a need for a unified, self-hosted platform combining "
    "comprehensive generation, deep customization, bulk processing, scanning, and history management.")

# P20 — Purpose (shorten)
set_text(doc.paragraphs[20],
    "QRCode Pro provides a unified platform for QR code workflows: 24+ QR types (URL, WiFi, vCard, UPI, WhatsApp, "
    "social media, etc.), 62 visual customization controls across 9 categories, bulk processing for up to 100 codes, "
    "integrated scanning via camera/upload, persistent history with favorites, and Chart.js dashboard analytics. "
    "All data stays local with bcrypt security and self-hosted deployment.")

# P24 — Scope intro
set_text(doc.paragraphs[24],
    "The scope of QRCode Pro includes the following key functionalities:")

# P26 — Scope bullet 1 (shorten)
set_text(doc.paragraphs[26],
    "QR Code Generation: 23 type-specific forms supporting URL, WiFi, vCard, UPI, WhatsApp, Email, "
    "SMS, Calendar, GPS, Bitcoin, YouTube, Twitter, Instagram, Facebook, LinkedIn, Spotify, Zoom, "
    "PDF, Image, MeCard, and App Store links with real-time preview.")

# P27 — Scope bullet 2 (shorten)
set_text(doc.paragraphs[27],
    "Customization Engine: 62 controls in 9 sections \u2014 color presets, dot style & gradient, "
    "background, eye frame, eye dot, logo/watermark, frame & label, shadow, and size/export settings.")

# P28 — Scope bullet 3 (shorten)
set_text(doc.paragraphs[28],
    "Export & Management: PNG/SVG/PDF download, print, clipboard copy, bulk ZIP export, "
    "searchable history with favorites, CSV export, and dashboard analytics with Chart.js charts.")

print("Fix 1: Shortened P1 content to prevent overflow \u2714")


# ═══════════════════════════════════════════════════════════════
# FIX 2: Remove extra empty paragraphs before footer (P29-31)
#         and after footer (P34-37) to save space
# ═══════════════════════════════════════════════════════════════
# P29, P30, P31 are empty — reduce to just P31 empty
set_text(doc.paragraphs[29], "")
set_text(doc.paragraphs[30], "")
# Keep P31 as single empty line before footer

# P34-37 are empty after footer — clear extras
set_text(doc.paragraphs[34], "")
set_text(doc.paragraphs[35], "")
set_text(doc.paragraphs[36], "")
set_text(doc.paragraphs[37], "")

# Reduce spacing on empty paras to minimize whitespace
for idx in [29, 30, 34, 35, 36, 37]:
    p = doc.paragraphs[idx]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if p.runs:
        p.runs[0].font.size = Pt(1)

print("Fix 2: Removed extra empty paragraphs \u2714")


# ═══════════════════════════════════════════════════════════════
# FIX 3: Fix T0 Row 0 Cell 0 — clean up leftover content
#         This merged header cell has 18 paragraphs, many are old
# ═══════════════════════════════════════════════════════════════
t0 = doc.tables[0]
cell0 = t0.rows[0].cells[0]
# para 0 = empty, para 1 = FACULTY header, para 2 = Department
# para 3+ = leftover content — replace with QRCode Pro scope details

cell_paras = cell0.paragraphs
# p3-p17 are leftover from password analyzer — rewrite them
CELL_CONTENT = [
    # p3: Scope item 1
    "QR Type Support: Generate QR codes for 24+ types including URL, WiFi, vCard, UPI, WhatsApp, Bitcoin, social media profiles, calendar events, GPS coordinates, and more.",
    # p4: Scope item 2
    "Visual Customization: 62 controls across 9 accordion sections — color presets, dot styles (6 options with gradients), background settings, eye frame/dot styling, logo overlay, frame & label, shadow effects, and export configuration.",
    # p5: Scope item 3
    "Bulk Generation: Generate up to 100 QR codes at once from manual text input or CSV file upload, with progress tracking and ZIP archive download including a manifest file.",
    # p6: empty
    "",
    # p7: Limitations heading
    "1.3 Limitations",
    # p8: empty
    "",
    # p9: Limitation 1
    "QR codes are generated entirely on the client side using JavaScript — the server stores only metadata (type, title, content, settings) not the actual QR images.",
    # p10: empty
    "",
    # p11: Limitation 2
    "Camera-based QR scanning requires HTTPS in production environments due to browser security policies for WebRTC access.",
    # p12: empty
    "",
    # p13: Limitation 3
    "SQLite database has limited concurrent write support, though this is mitigated by WAL journal mode. Not suitable for very high-traffic deployments.",
    # p14: empty
    "",
    # p15: Literature review heading
    "1.4 Existing System and Literature Review",
    # p16: empty
    "",
    # p17: Table label
    "Table:1 Existing System Study",
]

for pi, text in enumerate(CELL_CONTENT):
    actual_pi = pi + 3  # starts at p3
    if actual_pi < len(cell_paras):
        set_text(cell_paras[actual_pi], text)

print("Fix 3: Cleaned T0 merged header cell content \u2714")


# ═══════════════════════════════════════════════════════════════
# FIX 4: Fix T0 Row 6 — was showing "Developer tool only" in all cells
# ═══════════════════════════════════════════════════════════════
if len(t0.rows) >= 7:
    row6_data = [
        "GoQR.me API",
        "GoQR.me",
        "2023",
        "REST API for QR generation",
        "API-first, multiple output formats",
        "No UI, no auth, no customization beyond colors",
        "Java, REST API",
        "Developer tool only"
    ]
    for ci, val in enumerate(row6_data):
        if ci < len(t0.rows[6].cells):
            set_cell_text(t0.rows[6].cells[ci], val)

print("Fix 4: Fixed T0 row 6 literature entry \u2714")


# ═══════════════════════════════════════════════════════════════
# FIX 5: Fix Future Scope section (P39-71)
# ═══════════════════════════════════════════════════════════════

# P39 — Future Scope heading (already correct)
# P40 — intro text
set_text(doc.paragraphs[40],
    "The evolution of QRCode Pro will advance with emerging technologies and growing user demands. "
    "The future scope includes:")

# P42 — was: heading got too long, fix
set_text(doc.paragraphs[42], "AI-Powered QR Design")

# P44-46 — AI bullets (already correct from last edit)

# P48 — leftover heading "AI QR code complexity explanation" — fix
set_text(doc.paragraphs[48], "Dynamic QR Codes with Scan Analytics")

# P50-52 — Dynamic QR bullets (already correct)

# P53-57 — extra empty paras, collapse them
for idx in [53, 54, 55, 56, 57]:
    p = doc.paragraphs[idx]
    set_text(p, "")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if p.runs:
        p.runs[0].font.size = Pt(1)

# P58 — heading already good
# P60-61 — bullets already good

# P62 — leftover: "It will be very small and minimal application..." — fix
set_text(doc.paragraphs[62],
    "Offline-capable PWA mode for scanning QR codes without internet connection.")

# P63 — heading already good (API Marketplace)
# P65-66 — bullets already good

# P67 — leftover: "It will help users with timely info..." — fix
set_text(doc.paragraphs[67],
    "Automated QR code health monitoring with alerts for broken destination URLs.")

# P69 — heading already good (Enterprise Features)
# P71 — already good (White-label support)

print("Fix 5: Fixed Future Scope section \u2714")


# ═══════════════════════════════════════════════════════════════
# FIX 6: Fix Conclusion and References (P73-80)
# ═══════════════════════════════════════════════════════════════

# P73 — Conclusion heading (keep)
# P74 — conclusion text — rewrite
set_text(doc.paragraphs[74],
    "We identified the problem of fragmented QR code tools and defined a comprehensive solution \u2014 QRCode Pro. "
    "Through literature review, we analyzed existing systems and identified gaps that QRCode Pro addresses with "
    "24+ QR types, 62 customization controls, bulk generation, integrated scanning, and full history management.")

# P76 — References heading (keep)

# P77-80 — References (fix mangled URLs)
set_text(doc.paragraphs[77],
    "https://github.com/nickolay-bliokh/qr-code-styling \u2014 QR Code Styling Library (foundation for QR engine)")
set_text(doc.paragraphs[78],
    "https://www.qrmonkey.com \u2014 QR Monkey (existing system analyzed in literature review)")
set_text(doc.paragraphs[79],
    "https://goqr.me \u2014 GoQR.me API (existing system analyzed for comparison)")
set_text(doc.paragraphs[80],
    "https://davidshimjs.github.io/qrcodejs/ \u2014 QRCode.js Library (used for basic QR previews in history page)")

print("Fix 6: Fixed Conclusion and References \u2714")


# ═══════════════════════════════════════════════════════════════
# FIX 7: Check and fix T1 (second literature review table)
# ═══════════════════════════════════════════════════════════════
t1 = doc.tables[1]
if len(t1.rows) >= 7 and len(t1.columns) >= 8:
    # Fix row 0 cell 0 — merged header with leftover content
    t1_cell0 = t1.rows[0].cells[0]
    for pi, p in enumerate(t1_cell0.paragraphs):
        if pi >= 3:  # p0=empty, p1=FACULTY, p2=Department, p3+=content
            txt = p.text.strip()
            if txt and ("QR code" in txt.lower() or "breach" in txt.lower() or
                       "complexity" in txt.lower() or "brute" in txt.lower() or
                       "analyzer" in txt.lower() or "strength" in txt.lower()):
                # Leftover password content in T1 header cell — replace
                pass  # Will be handled by content mapping below

    # T1 rows 2-6 data (already set in edit_lab_manual.py, verify)
    T1_DATA = [
        ["QR Code Styling Library", "Nickolay Bliokh", "2023", "Client-side canvas rendering with dot and eye styling",
         "Deep customization, gradients, logos, shapes", "Library only, no full application", "JavaScript, Canvas API", "Foundation for QRCode Pro engine"],
        ["html5-qrcode Scanner", "mebjas", "2023", "Camera-based QR decoder using WebRTC",
         "Multi-camera support, fast decode, drag-drop", "Decode only, no generation capability", "JavaScript, WebRTC", "Used in QRCode Pro scanner module"],
        ["Chart.js Analytics", "Chart.js Team", "2023", "Canvas-based chart rendering library",
         "Responsive, many chart types, lightweight", "Client-side only, no server processing", "JavaScript, HTML5 Canvas", "Used in QRCode Pro dashboard"],
        ["JSZip for Bulk Export", "Stuart Knightley", "2023", "Client-side ZIP file creation",
         "No server needed, handles binary data", "Memory limits for very large archives", "JavaScript", "Used in bulk QR download feature"],
        ["jsPDF Document Export", "jsPDF Team", "2023", "Client-side PDF generation library",
         "No server needed, custom page layout", "Limited font support, canvas-based", "JavaScript, Canvas", "Used in QRCode Pro PDF export"],
    ]
    for ri, row_data in enumerate(T1_DATA):
        actual_row = ri + 2
        if actual_row < len(t1.rows):
            for ci, val in enumerate(row_data):
                if ci < len(t1.rows[actual_row].cells):
                    set_cell_text(t1.rows[actual_row].cells[ci], val)

    # Fix T1 header cell content (same structure as T0)
    t1_cell_paras = t1_cell0.paragraphs
    T1_CELL_CONTENT = [
        # p3
        "Integrated Scanner: Camera-based QR scanning using html5-qrcode library at 10fps with multi-camera support. File upload scanning with drag-and-drop. Auto-detection of content type (URL, email, phone, WiFi, text).",
        # p4
        "User Authentication: Registration with bcrypt password hashing, login via username or email, session regeneration to prevent fixation attacks, profile management with name/email update and password change.",
        # p5
        "Dashboard Analytics: Four stat cards (total codes, total scans, types used, plan), quick-generate grid for 8 popular types, recent codes display, Chart.js doughnut chart for type distribution and bar chart for 7-day activity.",
        # p6
        "",
        # p7
        "1.4 Limitations",
        # p8
        "",
        # p9
        "Maximum bulk generation is capped at 100 codes per batch to prevent browser memory issues during client-side rendering.",
        # p10
        "",
        # p11
        "SQLite database has limited concurrent write support. WAL journal mode mitigates this, but the system is not designed for high-concurrency enterprise deployments.",
        # p12
        "",
        # p13
        "QR code images are generated client-side and not stored on the server \u2014 only metadata is saved. If a user clears browser cache, they must regenerate the visual QR code.",
        # p14
        "",
        # p15
        "1.4 Existing System and Literature Review",
        # p16
        "",
        # p17
        "Table:2 Technology Study",
    ]
    for pi, text in enumerate(T1_CELL_CONTENT):
        actual_pi = pi + 3
        if actual_pi < len(t1_cell_paras):
            set_text(t1_cell_paras[actual_pi], text)

print("Fix 7: Fixed T1 second literature review table \u2714")


# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
doc.save(DOC_PATH)
print(f"\nAll Practical 1 fixes applied and saved!")
print(f"File: {DOC_PATH}")
